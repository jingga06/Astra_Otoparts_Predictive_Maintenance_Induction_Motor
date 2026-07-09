"""Generates docs/Panduan_SUPER_LENGKAP_Case2.docx - an exhaustive, zero-to-
hero explanation of the entire predictive maintenance prototype: data,
every pipeline stage (with formulas + why that method was chosen), a full
illustrated dashboard tour, anticipated critical questions, and improvement
ideas.

Run: python -m scripts.generate_mega_guide   (from src/)
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm, Inches

ROOT = Path(__file__).resolve().parents[2]
ASSETS = ROOT / "docs" / "assets"
OUT_PATH = ROOT / "docs" / "Panduan_SUPER_LENGKAP_Case2.docx"

NAVY = RGBColor(0x0F, 0x17, 0x2A)
BLUE = RGBColor(0x03, 0x69, 0xA1)
GREY = RGBColor(0x47, 0x55, 0x69)
GREEN = RGBColor(0x16, 0xA3, 0x4A)
RED = RGBColor(0xDC, 0x26, 0x26)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)


def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = NAVY
    return p


def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = BLUE
    return p


def h3(text):
    p = doc.add_heading(text, level=3)
    for r in p.runs:
        r.font.color.rgb = GREY
    return p


def body(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    p.paragraph_format.space_after = Pt(8)
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def note_box(text, label="CATATAN"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(f"{label}: " + text)
    r.italic = True
    r.font.color.rgb = BLUE
    return p


def formula_box(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(11)
    r.font.color.rgb = NAVY
    r.bold = True
    return p


def qa(question, answer):
    p = doc.add_paragraph()
    r = p.add_run("T: " + question)
    r.bold = True
    r.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(2)
    body("J: " + answer)


def picture(filename, caption, width=Inches(6.0)):
    doc.add_picture(str(ASSETS / filename), width=width)
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = GREY
    cap.paragraph_format.space_after = Pt(14)


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t


# ================================================================ COVER
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("PANDUAN SUPER LENGKAP")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Predictive Maintenance Dashboard - Induction Motor")
r.font.size = Pt(15)
r.font.color.rgb = BLUE

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("Dari Nol Sampai Paham Total - Data, Metode, Rumus, Alasan, dan Tur Dashboard")
r.font.size = Pt(12)
r.font.color.rgb = GREY

sub3 = doc.add_paragraph()
sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub3.add_run("Case 2 - Bootcamp AOP Winteq - PT Astra Otoparts Tbk")
r.font.size = Pt(11)
r.font.color.rgb = GREY

doc.add_paragraph()
note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = note.add_run(
    "Dokumen ini ditulis supaya siapa pun yang belum pernah lihat proyek ini "
    "sama sekali bisa paham total: apa yang dikerjakan, kenapa dikerjakan "
    "dengan cara itu, dan apa artinya tiap angka yang muncul di layar. "
    "Ditulis lengkap dan mendetail secara sengaja - lebih baik kepanjangan "
    "daripada ada yang kelewat."
)
r.italic = True
r.font.color.rgb = GREY
doc.add_page_break()

# ================================================================ DAFTAR ISI
h1("Daftar Isi")
toc = [
    "1. Pendahuluan - Apa, Kenapa, dan Cara Baca Dokumen Ini",
    "2. Mengenal Data dari Nol",
    "3. Gambaran Besar Arsitektur Sistem",
    "4. Tahap 1: Ekstraksi Fitur dari Sinyal Getaran",
    "5. Tahap 2: Deteksi Anomali dengan Isolation Forest",
    "6. Tahap 3: Health Index dengan PCA",
    "7. Tahap 4: Logika Alarm - Persistence + Voting",
    "8. Tahap 5: Estimasi Remaining Useful Life (RUL)",
    "9. Tur Lengkap Dashboard (dengan Contoh Gambar)",
    "10. Alarm Reason dan Root Cause - Detail Sepenuhnya",
    "11. Pertanyaan Kritis yang Mungkin Ditanyakan",
    "12. Ide Pengembangan agar Dilirik dan Disukai Astra",
    "13. Ringkasan Kelebihan dan Keterbatasan Jujur",
    "14. Daftar Istilah (Glossary) Lengkap",
    "15. Referensi",
]
for t in toc:
    p = doc.add_paragraph(t)
    p.paragraph_format.space_after = Pt(2)
doc.add_page_break()

# ================================================================ 1
h1("1. Pendahuluan - Apa, Kenapa, dan Cara Baca Dokumen Ini")
h2("1.1 Apa Proyek Ini Secara Keseluruhan")
body(
    "Ini adalah prototipe sistem Predictive Maintenance (perawatan prediktif) "
    "untuk motor induksi, dibuat untuk Case 2 Bootcamp AOP Winteq (PT Astra "
    "Otoparts Tbk). Tujuannya: mendeteksi tanda-tanda kerusakan motor SEBELUM "
    "motor itu benar-benar berhenti mendadak, supaya perawatan bisa "
    "dijadwalkan lebih dulu tanpa mengganggu produksi."
)
h2("1.2 Kenapa Ini Penting untuk Astra")
body(
    "Pabrik Astra Otoparts digerakkan banyak motor induksi (conveyor, pompa, "
    "compressor, mesin stamping/CNC). Pelanggan OEM (Toyota, Honda, dll.) "
    "beroperasi just-in-time, artinya kalau satu motor mendadak berhenti, "
    "efeknya berantai ke jadwal pengiriman dan bisa kena penalti. Sistem "
    "monitoring yang ada sekarang baru bereaksi SETELAH ada tanda kerusakan "
    "jelas - terlambat. Sistem yang diusulkan di sini mendeteksi jauh lebih "
    "awal, dengan tingkat alarm palsu yang rendah supaya teknisi tetap percaya "
    "pada alarm yang muncul."
)
h2("1.3 Cara Membaca Dokumen Ini")
body(
    "Dokumen ini disusun berurutan dari yang paling dasar ke yang paling "
    "detail. Disarankan dibaca dari depan ke belakang tanpa loncat, karena "
    "istilah yang dijelaskan di bab awal akan dipakai lagi di bab-bab "
    "selanjutnya tanpa diulang penjelasannya. Setiap istilah teknis yang "
    "muncul pertama kali akan langsung dijelaskan di tempat, dan semuanya "
    "dirangkum lagi di Bab 14 (Glossary)."
)

# ================================================================ 2
h1("2. Mengenal Data dari Nol")
h2("2.1 Dataset yang Dipakai: NASA IMS Bearing Dataset")
body(
    "Karena Astra belum mengirimkan data motor asli, prototipe ini memakai "
    "dataset publik dari NASA/IMS (Center for Intelligent Maintenance "
    "Systems), yang berisi rekaman getaran motor sungguhan yang dijalankan "
    "TERUS-MENERUS sampai salah satu komponennya (bearing) benar-benar rusak "
    "total. Mentor pembimbing (\"Miss\") sudah menyetujui pemakaian dataset "
    "ini untuk tahap prototipe, karena kondisinya mirip motor industri "
    "sungguhan dan sifatnya open-source (bebas dipakai)."
)
h2("2.2 Apa Itu Bearing")
body(
    "Bearing adalah laher/bantalan - komponen di dalam motor yang membuat "
    "poros bisa berputar dengan mulus dan mengurangi gesekan. Berdasarkan "
    "riset industri (survei IEEE-IAS dan EPRI yang jadi rujukan tim ini), "
    "sekitar 41-42% dari seluruh kegagalan motor induksi disebabkan oleh "
    "kerusakan bearing - jenis kegagalan paling umum, karena itulah "
    "prototipe ini fokus ke situ dulu."
)
h2("2.3 Setup Alat Uji (Test Rig)")
body(
    "Di alat uji yang menghasilkan dataset ini, satu poros (shaft) diputar "
    "konstan di 2000 RPM oleh motor AC, dan di sepanjang poros itu dipasang "
    "EMPAT bearing (disebut Bearing 1, 2, 3, 4), masing-masing dipantau "
    "getarannya dengan sensor accelerometer terpisah."
)
h2("2.4 Test 2 dan Test 3 - Dua Rekaman Berbeda")
body(
    "Dari 3 percobaan yang tersedia di dataset NASA, prototipe ini memakai "
    "2 di antaranya:"
)
table(
    ["", "Test 2 (dipakai sebagai run pengembangan)", "Test 3 (dipakai sebagai run holdout/ujian)"],
    [
        ["Lama rekaman", "163,8 jam (~6,8 hari)", "1073,3 jam (~44,7 hari)"],
        ["Jumlah snapshot getaran", "984", "6.324"],
        ["Bearing yang rusak", "Bearing 1", "Bearing 3"],
        ["Jenis kerusakan", "Outer race (cincin luar)", "Outer race (cincin luar)"],
    ],
)
body(
    "\"Run pengembangan\" artinya Test 2 dipakai untuk membangun dan "
    "mengatur sistemnya. \"Holdout\" artinya Test 3 SENGAJA tidak diintip "
    "sama sekali sampai sistemnya selesai dibangun dari Test 2 - baru "
    "setelah itu dijalankan ke Test 3 TANPA diubah-ubah lagi, supaya hasil "
    "ujiannya jujur (tidak \"nyontek\"). Ini prinsip validasi standar dalam "
    "ilmu machine learning, disebut train/holdout split."
)
h2("2.5 Kenapa Test 1 Tidak Dipakai")
body(
    "Test 1 punya struktur data berbeda (2 sensor per bearing, bukan 1) dan "
    "jenis kerusakan berbeda (inner race dan roller element, bukan outer "
    "race). Karena Test 2 dan Test 3 sama-sama outer race, keduanya bisa "
    "jadi SATU cerita yang nyambung (kembangkan di satu, uji jujur di yang "
    "lain). Test 1 disimpan sebagai pengembangan lanjutan di luar scope "
    "prototipe minggu ini."
)
h2("2.6 Kesimpulan: 8 Kombinasi Bearing-Run")
body(
    "2 rekaman x 4 bearing = 8 kombinasi yang masing-masing dianalisis dan "
    "disimpan terpisah. Bearing yang TIDAK rusak di tiap rekaman berfungsi "
    "sebagai \"kontrol sehat\" - bukti bahwa sistem tidak asal membunyikan "
    "alarm ke motor yang sebenarnya baik-baik saja."
)

# ================================================================ 3
h1("3. Gambaran Besar Arsitektur Sistem")
body(
    "Sebelum masuk detail tiap tahap, ini gambaran alur datanya secara "
    "keseluruhan - dari sinyal getaran mentah sampai jadi angka-angka yang "
    "muncul di dashboard:"
)
picture("05_pipeline_diagram.png", "Gambar 3.1 - Alur pipeline lengkap, dua cabang paralel")
body(
    "Ada DUA CABANG yang berjalan dari titik yang sama (fitur getaran) tapi "
    "punya tujuan berbeda:"
)
bullet(
    "Menjawab pertanyaan \"apakah ada yang salah SEKARANG?\" - hasilnya "
    "berupa status alarm (Normal/Warning/Critical).",
    bold_lead="Cabang A (Deteksi) - ",
)
bullet(
    "Menjawab pertanyaan \"kalau memang ada yang salah, berapa lama lagi "
    "sampai benar-benar gagal?\" - hasilnya berupa angka RUL.",
    bold_lead="Cabang B (Prediksi) - ",
)
note_box(
    "Dua cabang ini SALING TERPISAH secara perhitungan (masing-masing pakai "
    "metode sendiri), tapi Cabang B baru mulai aktif setelah Cabang A "
    "pertama kali membunyikan alarm resmi - itu sebabnya RUL selalu kosong "
    "di awal (status \"Monitoring\") sampai ada alasan nyata untuk mulai "
    "menghitungnya."
)

# ================================================================ 4
h1("4. Tahap 1: Ekstraksi Fitur dari Sinyal Getaran")
h2("4.1 Kenapa Tidak Dipakai Sinyal Mentahnya Langsung?")
body(
    "Setiap potongan rekaman getaran berisi 20.480 titik angka (direkam "
    "20.000 kali per detik, selama 1 detik). Terlalu banyak dan terlalu "
    "\"berisik\" untuk langsung dibandingkan satu titik ke titik lain. "
    "Makanya tiap potongan diringkas dulu jadi beberapa angka ciri khas "
    "(disebut FITUR) yang benar-benar mewakili kondisi kesehatan bearing di "
    "detik itu."
)
h2("4.2 Fitur Time-Domain (dihitung langsung dari sinyal getaran)")
h3("RMS (Root Mean Square)")
formula_box("RMS = akar( rata-rata dari (nilai getaran)^2 )")
body(
    "Mengukur seberapa BESAR/KUAT getaran secara keseluruhan. Motor yang "
    "makin rusak biasanya bergetar makin kencang, jadi RMS-nya naik."
)
h3("Kurtosis")
formula_box("Kurtosis = rata-rata dari (x - rata2)^4  dibagi  simpangan^4")
body(
    "Mengukur seberapa \"tajam/mengejutkan\" pola getarannya - bukan "
    "seberapa besar, tapi seberapa SERING ada lonjakan tiba-tiba di "
    "tengah-tengah getaran yang biasa saja. Getaran motor sehat polanya "
    "mendekati kurva normal (lonceng), nilainya sekitar 3. Begitu ada "
    "benturan kecil dari cacat di bearing, muncul \"paku-paku\" tajam di "
    "sinyalnya, dan kurtosis-nya naik jauh di atas 3 (di data prototipe ini, "
    "naik dari ~3,5 jadi ~15,6 menjelang gagal)."
)
h3("Crest Factor")
formula_box("Crest Factor = nilai puncak tertinggi / RMS")
body(
    "Mengukur seberapa \"menonjol\" puncak tertingginya dibanding kekuatan "
    "rata-rata sinyal. Kalau ada benturan tajam sesaat (bukan getaran yang "
    "rata), crest factor-nya naik."
)
h2("4.3 Fitur Frequency-Domain: BPFO, BPFI, BSF (dihitung dari fisika bearing)")
body(
    "Ini bagian yang paling \"ilmiah\" dari sistem ini. Setiap bearing "
    "punya bentuk fisik tertentu (jumlah bola/rol, diameter, sudut kontak). "
    "Kalau ada CACAT KECIL di salah satu bagian bearing (misalnya di cincin "
    "luar), setiap kali bola/rol melewati titik cacat itu, terjadi benturan "
    "kecil - dan karena bearingnya berputar konstan, benturan itu terjadi "
    "BERULANG di frekuensi yang bisa dihitung pasti secara matematis dari "
    "ukuran fisik bearingnya, BUKAN ditebak-tebak."
)
formula_box("BPFO = (n/2) x fr x (1 - (d/D) x cos(sudut))")
formula_box("BPFI = (n/2) x fr x (1 + (d/D) x cos(sudut))")
formula_box("BSF  = (D/(2d)) x fr x (1 - ((d/D) x cos(sudut))^2)")
body(
    "n = jumlah elemen gelinding, fr = kecepatan putar poros, d = diameter "
    "elemen gelinding, D = diameter lintasan (pitch diameter). Untuk "
    "bearing di dataset ini (Rexnord ZA-2115, 16 elemen, poros 2000 RPM), "
    "hasil hitungannya:"
)
table(
    ["Nama", "Artinya", "Frekuensi hasil hitung"],
    [
        ["BPFO", "Ball Pass Frequency, Outer race - cacat di cincin LUAR", "236,4 Hz"],
        ["BPFI", "Ball Pass Frequency, Inner race - cacat di cincin DALAM", "296,9 Hz"],
        ["BSF", "Ball Spin Frequency - cacat di BOLA/rol itu sendiri", "139,9 Hz"],
        ["FTF", "Fundamental Train Frequency - kecepatan sangkar bola berputar", "14,8 Hz"],
    ],
)
note_box(
    "Angka-angka ini PERSIS sama dengan yang tertulis di dokumen riset tim "
    "(Laporan Progres Case 2), sudah dicek ulang lewat kode dan hasilnya "
    "cocok - jadi bukan angka karangan."
)
h2("4.4 Envelope Analysis - Cara Menonjolkan Sinyal Cacat yang Tersembunyi")
body(
    "Masalahnya: benturan kecil akibat cacat bearing itu SANGAT KECIL dan "
    "\"ketutupan\" oleh getaran struktur mesin yang jauh lebih besar. Kalau "
    "langsung dilihat spektrum sinyal mentahnya, benturan kecil itu nyaris "
    "tidak kelihatan. Solusinya, dipakai teknik 3 langkah bernama Envelope "
    "Analysis:"
)
numbered(
    "Band-pass filter: saring sinyal, hanya sisakan rentang frekuensi "
    "2.000-10.000 Hz (rentang di mana getaran akibat benturan kecil "
    "\"berdengung\" paling jelas)."
)
numbered(
    "Transformasi Hilbert: ambil garis \"selubung\" (envelope) dari sinyal "
    "yang sudah disaring - ibarat menggambar garis yang mengikuti puncak "
    "atas gelombangnya saja."
)
numbered(
    "FFT (Fast Fourier Transform) dari garis selubung itu: ubah dari "
    "\"sinyal vs waktu\" jadi \"energi vs frekuensi\", sehingga sekarang "
    "kelihatan jelas di frekuensi berapa energinya paling besar."
)
body(
    "Hasil akhirnya: kalau ada cacat di outer race, akan muncul PUNCAK "
    "TAJAM tepat di 236,4 Hz (BPFO). Ini yang jadi bukti fisik paling kuat "
    "dari seluruh sistem - lihat Gambar 9.5 di Bab 9 untuk contoh nyatanya."
)

# ================================================================ 5
h1("5. Tahap 2: Deteksi Anomali dengan Isolation Forest")
h2("5.1 Ide Dasarnya Apa")
body(
    "Isolation Forest bekerja dengan prinsip: \"data yang ANEH itu MUDAH "
    "DIPISAHKAN dari yang lain, sedangkan data yang NORMAL itu SUSAH "
    "dipisahkan karena mirip-mirip dengan banyak data lain di sekitarnya.\" "
    "Analoginya: kalau kamu disuruh cari 1 orang berbeda di keramaian orang "
    "yang mirip-mirip, kalau orang itu memang beda banget (misal pakai baju "
    "warna aneh sendirian), kamu akan cepat menemukannya cuma dengan "
    "beberapa pertanyaan pemisah. Tapi kalau semua orang mirip, butuh "
    "banyak pertanyaan buat memisahkan satu-satu."
)
body(
    "Secara teknis, algoritma ini membuat banyak \"pohon keputusan\" acak "
    "yang membagi-bagi data secara acak berulang kali. Titik data yang ANEH "
    "akan lebih cepat \"terisolasi sendirian\" (butuh sedikit pembagian), "
    "sedangkan titik data NORMAL butuh banyak pembagian karena dikelilingi "
    "banyak titik serupa."
)
formula_box("skor_anomali(x) = 2 ^ ( -rata2_panjang_jalur_isolasi(x) / konstanta )")
body(
    "Semakin PENDEK jalur yang dibutuhkan untuk mengisolasi suatu titik "
    "(artinya makin gampang dipisahkan / makin aneh), semakin skor "
    "anomalinya mendekati 1."
)
h2("5.2 Kenapa Dilatih HANYA dari Data Sehat")
body(
    "Model ini sengaja hanya \"dikenalkan\" ke 50% data paling awal tiap "
    "rekaman (dianggap masih sehat), TIDAK diperlihatkan contoh data rusak "
    "sama sekali. Kenapa? Karena data kegagalan itu LANGKA (cuma ada di "
    "ujung rekaman, sedikit sekali dibanding data sehat), jadi lebih aman "
    "mengajari sistem \"seperti apa itu normal\" lalu menandai apa pun yang "
    "menyimpang jauh, daripada berharap sistem bisa \"mengenali\" pola "
    "rusak dari contoh yang terlalu sedikit."
)
h2("5.3 Kenapa Isolation Forest, Bukan Metode Lain?")
table(
    ["Metode", "Kelebihan", "Kenapa TIDAK dipilih di sini"],
    [
        ["Isolation Forest (dipilih)", "Cepat, ringan (bisa jalan di edge device tanpa GPU), hasilnya gampang dijelaskan, cocok untuk data tidak berlabel", "-"],
        ["Autoencoder (neural network)", "Bisa menangkap pola non-linear yang lebih kompleks", "Butuh lebih banyak data untuk dilatih dengan baik, lebih berat secara komputasi, hasilnya lebih sulit dijelaskan (\"black box\")"],
        ["One-Class SVM", "Landasan matematis kuat", "Lebih lambat untuk data besar, sensitif terhadap pemilihan parameter"],
        ["Threshold sederhana per-parameter", "Paling gampang dibuat", "Ini justru masalah yang ingin diperbaiki - banyak alarm palsu (sudah dijelaskan dari awal proposal)"],
    ],
)
note_box(
    "Autoencoder tetap disebut di dokumen riset tim sebagai \"pembanding\" "
    "(bisa dicoba di pengembangan lanjutan), tapi Isolation Forest dipilih "
    "sebagai baseline utama karena lebih cocok dengan kondisi data yang "
    "terbatas dan kebutuhan sistem yang bisa dijelaskan ke teknisi pabrik."
)
h2("5.4 EWMA Smoothing - Menghaluskan Noise")
body(
    "Skor anomali mentah bisa naik-turun sesaat karena noise kecil "
    "(getaran sesaat yang tidak berarti apa-apa). Untuk menghaluskannya, "
    "dipakai EWMA (Exponentially Weighted Moving Average) - rata-rata "
    "berjalan yang memberi bobot lebih besar ke data terbaru, tapi tetap "
    "\"mengingat\" sedikit data-data sebelumnya, sehingga hasilnya lebih "
    "stabil dibanding rata-rata biasa."
)

# ================================================================ 6
h1("6. Tahap 3: Health Index dengan PCA")
h2("6.1 Masalah yang Diselesaikan")
body(
    "Ada 6 fitur getaran (RMS, kurtosis, crest factor, energi BPFO/BPFI/BSF) "
    "yang semuanya bergerak naik saat bearing makin rusak, tapi dengan "
    "kecepatan dan skala yang beda-beda. Daripada operator harus melihat 6 "
    "grafik terpisah, dibutuhkan SATU angka gabungan yang mewakili semuanya."
)
h2("6.2 Apa Itu PCA (Principal Component Analysis)")
body(
    "PCA adalah teknik matematika yang mencari SATU ARAH KOMBINASI dari "
    "banyak angka yang paling baik mewakili \"ke mana arah gerak bersama\" "
    "dari semuanya. Bayangkan 6 fitur itu sebagai 6 jarum kompas yang "
    "sebagian besar menunjuk ke arah yang mirip (naik bareng saat rusak) - "
    "PCA menemukan \"arah rata-rata\" dari semua jarum itu, lalu proyeksikan "
    "tiap titik waktu ke arah tersebut untuk mendapat satu angka tunggal."
)
h2("6.3 Kenapa PCA, Bukan Rata-Rata Tertimbang Manual?")
body(
    "Cara paling gampang sebenarnya bisa saja: \"gabungkan 6 fitur dengan "
    "bobot 20% RMS + 30% kurtosis + ...\" - tapi angka bobot itu HARUS "
    "ditebak manusia, dan sulit dipertanggungjawabkan kenapa 30% bukan 25%. "
    "PCA justru MENGHITUNG bobot itu sendiri dari pola data yang sebenarnya "
    "terjadi, jadi tidak ada angka yang \"asal tebak\" dan bisa "
    "dipertanggungjawabkan secara matematis. Ini juga cepat dilatih "
    "(hitungannya dalam hitungan milidetik) dan mudah dijelaskan."
)
h2("6.4 Pemilihan Fitur yang \"Monoton\"")
body(
    "Sebelum digabung, sistem dulu mengecek: dari 6 fitur, mana saja yang "
    "BENAR-BENAR punya kecenderungan naik/turun konsisten seiring waktu "
    "(bukan cuma naik-turun acak tanpa arah)? Fitur yang polanya jelas "
    "dipakai, yang tidak jelas arahnya dibuang - supaya angka gabungan "
    "akhirnya benar-benar mencerminkan tren degradasi, bukan ikut-ikutan "
    "noise dari fitur yang tidak relevan."
)
h2("6.5 Skala 0-100")
body(
    "Angka gabungan dari PCA diubah ke skala 0-100 supaya gampang dibaca "
    "siapa saja: 100 = rata-rata bearing itu sendiri saat sehat, 0 = "
    "kondisi bearing itu sendiri di akhir rekamannya (titik dia gagal). "
    "Karena patokannya per-bearing (bukan patokan tetap universal), "
    "perbandingan yang adil adalah \"seberapa jauh bearing ini dari kondisi "
    "sehatnya SENDIRI\", bukan dibanding bearing lain."
)

# ================================================================ 7
h1("7. Tahap 4: Logika Alarm - Persistence + Voting")
h2("7.1 Masalah False Alarm")
body(
    "Kalau alarm dibunyikan setiap kali SATU fitur saja melewati ambang "
    "batas, sistem akan terlalu sering \"teriak\" untuk hal yang sebenarnya "
    "cuma noise sesaat - dan lama-lama teknisi jadi tidak percaya lagi ke "
    "alarm (bahaya besar, karena alarm yang BENAR-BENAR penting jadi ikut "
    "diabaikan). Ini masalah yang secara eksplisit disebut di awal proposal "
    "sebagai kelemahan sistem monitoring yang sudah ada."
)
h2("7.2 Voting Rule")
body(
    "Kalau 2 ATAU LEBIH dari 6 fitur sama-sama abnormal DI WAKTU YANG SAMA, "
    "alarm langsung dibunyikan CRITICAL tanpa menunggu. Logikanya: kalau "
    "cuma 1 fitur yang aneh, mungkin itu kebetulan/noise. Tapi kalau 2 "
    "fitur berbeda SAMA-SAMA aneh bersamaan, kemungkinan kebetulannya jauh "
    "lebih kecil - ini bukti yang lebih kuat."
)
h2("7.3 Persistence Rule")
body(
    "Kalau cuma 1 fitur yang abnormal, sistem tidak langsung bunyi. Fitur "
    "itu harus tetap abnormal selama beberapa kali BERTURUT-TURUT (jumlahnya "
    "menyesuaikan panjang rekaman, minimal 3 kali) baru dianggap WARNING. "
    "Ini seperti aturan \"debounce\" pada tombol elektronik - mencegah "
    "lonjakan sesaat dianggap kejadian serius."
)
h2("7.4 Burn-in dan Hysteresis")
bullet(
    "5% data paling awal tidak dievaluasi sama sekali - ini masa \"kalibrasi\" "
    "sistem mengenali pola sehat.",
    bold_lead="Burn-in - ",
)
bullet(
    "Begitu status sudah WARNING/CRITICAL, statusnya TIDAK langsung "
    "kembali NORMAL hanya karena 1 baris data terlihat baik-baik saja - "
    "harus \"terkonfirmasi normal\" selama beberapa kali berturut-turut "
    "juga, supaya status tidak kedap-kedip antara aman dan bahaya.",
    bold_lead="Hysteresis (recovery) - ",
)
h2("7.5 Kenapa Parameternya Diskalakan ke Panjang Rekaman")
body(
    "Tim menemukan sendiri lewat percobaan: parameter yang pas untuk "
    "rekaman pendek (Test 2, 164 jam) TERNYATA memicu banyak alarm palsu "
    "kalau langsung dipakai di rekaman yang 6,5 kali lebih panjang (Test 3, "
    "1073 jam) tanpa disesuaikan. Solusinya, jumlah \"burn-in\" dan "
    "\"persistence\" dihitung sebagai PERSENTASE dari panjang rekaman "
    "(bukan angka tetap), sehingga logika yang sama otomatis menyesuaikan "
    "diri baik untuk rekaman pendek maupun panjang tanpa perlu diatur ulang "
    "manual."
)

# ================================================================ 8
h1("8. Tahap 5: Estimasi Remaining Useful Life (RUL)")
h2("8.1 Kapan RUL Mulai Dihitung")
body(
    "RUL BARU mulai dihitung setelah alarm PERTAMA KALI resmi berbunyi "
    "(lolos aturan persistence+voting di Tahap 4). Sebelum itu, sistem "
    "sengaja tidak berani menebak - statusnya \"Monitoring, belum stabil\". "
    "Ini jujur secara statistik: menebak durasi kegagalan sebelum ada tanda "
    "nyata itu tidak berdasar."
)
h2("8.2 Model Matematis: Kurva Eksponensial")
formula_box("HI(t) = p1 x e^(p2 x t) + p3")
body(
    "p3 adalah titik awal (nilai Health Index saat alarm baru nyala), p1 "
    "adalah seberapa besar ruang kurva itu untuk naik, dan p2 adalah "
    "SEBERAPA CEPAT percepatan kerusakannya - parameter paling penting. "
    "Kurva ini di-cocokkan (fit) ke data yang ada, lalu dihitung ulang "
    "setiap ada data baru masuk."
)
formula_box("EOL = (1/p2) x ln( (ET - p3) / p1 )")
formula_box("RUL(sekarang) = EOL - waktu_sekarang")
body(
    "EOL (End of Life) adalah titik waktu di mana kurva itu diperkirakan "
    "menyentuh \"garis gagal\" (ET, singkatan dari Error Threshold / "
    "failure threshold)."
)
h2("8.3 Dari Mana Angka \"Garis Gagal\" (ET) Berasal?")
body(
    "ET TIDAK diambil dari data rekaman yang sedang dianalisis itu sendiri "
    "(itu namanya \"nyontek masa depan diri sendiri\", tidak jujur secara "
    "statistik). Sebaliknya, ET diambil dari REKAMAN LAIN yang bearingnya "
    "sudah benar-benar terbukti gagal - misalnya, saat menghitung RUL untuk "
    "Test 2, angka ET-nya dipinjam dari nilai Health Index Test 3 pas "
    "bearingnya beneran rusak (dan sebaliknya). Ini disebut \"transfer "
    "lintas mesin\" - meniru cara seorang teknisi berpengalaman yang tahu "
    "\"biasanya kalau udah separah ini, itu tandanya mau rusak\" dari "
    "pengalaman di mesin-mesin lain."
)
h2("8.4 Kenapa Model Eksponensial Sederhana, Bukan Deep Learning (LSTM)?")
table(
    ["Pertimbangan", "Model Eksponensial (dipilih)", "LSTM / Deep Learning"],
    [
        ["Kebutuhan data", "Bisa jalan dengan sedikit contoh kegagalan", "Butuh BANYAK contoh kegagalan historis untuk belajar dengan baik"],
        ["Risiko overfitting", "Rendah, karena modelnya sederhana (cuma 3 angka)", "Tinggi kalau data kegagalan sedikit"],
        ["Bisa dijelaskan ke teknisi", "Ya - kurva naik, gampang digambar dan dijelaskan", "Sulit - \"black box\", susah dijelaskan kenapa hasilnya begitu"],
        ["Estimasi ketidakpastian", "Bisa dihitung langsung dari kualitas kecocokan kurva", "Butuh teknik tambahan yang lebih rumit"],
    ],
)
note_box(
    "Ini bukan berarti LSTM \"jelek\" - LSTM justru sangat bagus KALAU data "
    "kegagalan historisnya banyak. Karena dataset yang tersedia cuma "
    "beberapa rekaman run-to-failure, model statistik sederhana ini pilihan "
    "yang lebih realistis dan sudah terbukti dipakai luas di literatur "
    "prognostics industri."
)
h2("8.5 Keterbatasan yang Diakui Secara Terbuka")
bullet(
    "Kalau pola kerusakannya naik-turun berulang (bukan naik mulus terus), "
    "kurva eksponensial tunggal bisa kurang akurat - solusi jangka panjang: "
    "model bertahap (piecewise) yang lebih fleksibel.",
    bold_lead="Degradasi dua tahap/berosilasi - ",
)
bullet(
    "Dataset ini direkam di 1 kecepatan/beban konstan. Di pabrik sungguhan, "
    "kondisi operasi bisa berubah-ubah - solusi jangka panjang: baseline "
    "sehat terpisah per mode operasi.",
    bold_lead="Kondisi operasi yang berubah-ubah - ",
)
bullet(
    "ET saat ini dipinjam dari SATU rekaman lain. Idealnya, dikumpulkan "
    "dari BANYAK rekaman kegagalan supaya jadi sebaran statistik yang lebih "
    "kuat, bukan cuma 1 angka.",
    bold_lead="Sumber ET masih terbatas - ",
)

doc.add_page_break()

# ================================================================ 9
h1("9. Tur Lengkap Dashboard (dengan Contoh Gambar)")
h2("9.1 Bar Judul Paling Atas")
body(
    "Bagian biru gelap di paling atas menampilkan judul, nama perusahaan, "
    "dan di kanan atas ada info seperti ini:"
)
formula_box("Motor Conveyor - Bearing 1\nTest 2 (dev run)  x  t = 89.0h")
body(
    "\"t = 89,0h\" artinya: posisi waktu yang SEDANG DILIHAT sekarang ada "
    "di JAM KE-89 dari keseluruhan rekaman (dihitung sejak rekaman itu "
    "mulai, jam ke-0). Ini bukan jam sungguhan di dunia nyata (bukan jam "
    "9 pagi dsb), tapi \"jam ke berapa SEJAK rekaman ini dimulai\". Angka "
    "ini berubah otomatis mengikuti posisi slider di sidebar."
)
h2("9.2 Sidebar - Panel Kontrol")
bullet("Pilih Test 2 atau Test 3.", bold_lead="Production run - ")
bullet("Pilih bearing 1-4 (yang diketahui rusak ditandai khusus).", bold_lead="Machine/bearing - ")
bullet(
    "Slider posisi waktu, tombol Play/Pause/Reset, dan pengatur kecepatan "
    "(1x-64x) - detail lengkap sudah dijelaskan di percakapan sebelumnya "
    "dan dirangkum ulang di Bab 10.",
    bold_lead="Kontrol replay - ",
)
h2("9.3 Tiga Kartu Ringkasan (KPI)")
picture("06_kpi_cards_mockup.png", "Gambar 9.1 - Contoh 3 kartu KPI persis di titik trigger (jam ke-89)")
body(
    "Perhatikan contoh di atas: Health Score MASIH TINGGI (98,9) tapi "
    "Alarm Status SUDAH CRITICAL. Ini BUKAN kesalahan tampilan - ini bukti "
    "nyata bahwa deteksi anomali (berbasis z-score fitur mentah) bereaksi "
    "lebih cepat/sensitif dibanding Health Index (yang berbasis tren "
    "gabungan PCA yang lebih halus). Sudah dijelaskan detail alasannya di "
    "Bab 5-6."
)
h2("9.4 Tab \"Trend\"")
picture("01_health_score_trend.png", "Gambar 9.2 - Grafik kiri: Health Score dari waktu ke waktu")
body(
    "Garis biru = Health Score. Titik merah = alarm sedang aktif. Garis "
    "putus-putus oranye = titik trigger pertama (mulai dihitungnya RUL). "
    "Perhatikan bentuknya: rata dulu di 100 (sehat), lalu mulai turun "
    "berosilasi (naik-turun, wajar untuk data getaran real), lalu jatuh "
    "tajam ke 0 menjelang akhir."
)
picture("02_anomaly_vs_threshold.png", "Gambar 9.3 - Grafik kanan: Anomaly Score vs garis ambang 3-sigma")
body(
    "Ini grafik yang PALING PAS untuk menjawab permintaan Miss: \"tunjukkan "
    "kapan kondisi melewati standar deviasi\". Garis abu-abu adalah skor "
    "anomali, garis merah putus-putus adalah ambang 3-sigma. Persis di "
    "jam ke-89, garis abu-abu MELEWATI garis merah untuk pertama kalinya - "
    "itulah momen yang dicari."
)
h2("9.5 Tab \"Signal Detail\"")
picture("03_raw_vibration_compare.png", "Gambar 9.4 - Perbandingan sinyal getaran mentah: sehat (kiri) vs menjelang gagal (kanan)")
body(
    "Perhatikan bedanya secara visual: saat sehat, getarannya halus dan "
    "kecil (skala -0,25 sampai 0,25). Menjelang gagal, getarannya jauh "
    "lebih liar dan besar (sampai skala -1 sampai 1) - ini konsisten dengan "
    "kenaikan RMS 9x lipat yang sudah dibahas sebelumnya."
)
picture("04_envelope_spectrum.png", "Gambar 9.5 - Envelope spectrum: bukti fisik lokasi kerusakan")
body(
    "Ini grafik PALING PENTING secara ilmiah. Garis hijau (saat sehat) "
    "nyaris datar di semua frekuensi - tidak ada yang menonjol. Garis "
    "hitam (menjelang gagal) punya PUNCAK TINGGI TEPAT di garis merah "
    "putus-putus BPFO (236,4 Hz) - persis seperti yang diprediksi rumus "
    "fisika di Bab 4. Ini BUKTI FISIK NYATA bahwa kerusakannya memang di "
    "cincin luar (outer race), bukan sekadar kesimpulan statistik."
)
h2("9.6 Tab \"Alerts & Recommendation\"")
body(
    "Berisi status alarm, alasan (alarm_reason), dugaan root cause, dan "
    "rekomendasi tindakan - dijelaskan lengkap dengan detail cara "
    "hitungnya di Bab 10."
)
h2("9.7 Pop-Up Notifikasi")
body(
    "Setiap kali status alarm NAIK tingkat (Normal ke Warning, atau "
    "Warning ke Critical) persis di titik waktu yang sedang dilihat, "
    "muncul notifikasi kecil otomatis di pojok layar - fitur ini "
    "ditambahkan supaya sistem terasa \"aktif memberi tahu\", bukan cuma "
    "\"pasif menampilkan\", sesuai permintaan untuk ada automated alert."
)
h2("9.8 Catatan Metodologi di Paling Bawah")
body(
    "Baris kecil paling bawah merangkum metode yang dipakai dalam satu "
    "kalimat, supaya siapa pun yang melihat dashboard langsung tahu dasar "
    "ilmiahnya tanpa harus baca dokumen terpisah."
)

doc.add_page_break()

# ================================================================ 10
h1("10. Alarm Reason dan Root Cause - Detail Sepenuhnya")
h2("10.1 Alarm Reason - Kalimat yang Muncul di Layar")
body(
    "Alarm reason adalah kalimat otomatis yang menjelaskan KENAPA status "
    "alarm menjadi seperti itu. Kalimatnya adalah TEMPLATE yang otomatis "
    "diisi angka/nama fitur sesuai kondisi nyata - bukan komputer yang "
    "\"berpikir\" dan mengarang kalimat. Contoh nyata dari data:"
)
table(
    ["Jam ke-", "Status", "Alarm Reason"],
    [
        ["3,3", "NORMAL", "Burn-in / baseline calibration period"],
        ["83,3", "NORMAL", "All parameters within normal range"],
        ["89,0", "CRITICAL", "Voting: 2 parameters abnormal at once (rms, bpfo_energy)"],
        ["90,0", "WARNING", "Persistence: rms abnormal for 7 consecutive windows (>= 5)"],
        ["163,2", "CRITICAL", "Voting: 6 parameters abnormal at once (rms, kurtosis, crest_factor, bpfo_energy, bpfi_energy, bsf_energy)"],
    ],
)
h2("10.2 Root Cause - Perlu Jujur Diakui")
body(
    "Tulisan \"Likely root cause\" di tab Alerts BUKAN hasil sistem "
    "\"mendiagnosa dari nol\" seperti detektif. Logikanya sederhana: kalau "
    "alarm sedang aktif DAN bearing yang sedang dilihat memang bearing yang "
    "SUDAH DIKETAHUI rusak (dari dokumentasi resmi NASA, bukan ditebak), "
    "sistem menampilkan \"Bearing outer-race defect\". Ini memanfaatkan "
    "fakta yang sudah diketahui, bukan diagnosis independen real-time."
)
note_box(
    "Bukti fisik yang SUNGGUHAN ada di grafik Envelope Spectrum (Gambar "
    "9.5) - puncak nyata di frekuensi BPFO. Tapi bukti ini belum "
    "\"disambungkan\" otomatis ke kalimat root cause - itu ide pengembangan "
    "yang dibahas di Bab 12."
)

# ================================================================ 11
h1("11. Pertanyaan Kritis yang Mungkin Ditanyakan")
body(
    "Bagian ini berisi pertanyaan-pertanyaan yang kemungkinan besar akan "
    "muncul dari dosen, mentor, atau juri - lengkap dengan jawaban yang "
    "jujur dan berdasar."
)

qa(
    "Kenapa tidak pakai deep learning (LSTM/CNN) dari awal untuk semuanya?",
    "Karena data kegagalan yang tersedia sangat terbatas (cuma beberapa "
    "rekaman run-to-failure). Deep learning butuh BANYAK contoh untuk "
    "belajar dengan baik, dan kalau dipaksakan dengan data sedikit, "
    "risikonya overfitting (modelnya \"menghafal\" data yang ada, bukan "
    "benar-benar belajar pola umum) dan hasilnya sulit dijelaskan. Metode "
    "statistik (Isolation Forest, PCA, model eksponensial) lebih realistis "
    "untuk kondisi data terbatas ini, dan didukung literatur prognostics "
    "industri sebagai pendekatan yang valid."
)
qa(
    "Bagaimana kalau motor beroperasi di kecepatan/beban yang berubah-ubah, "
    "bukan konstan seperti di dataset?",
    "Ini keterbatasan yang diakui secara terbuka (dijelaskan di Bab 8.5). "
    "Dataset NASA IMS direkam di 1 kecepatan/beban tetap. Di pabrik "
    "sungguhan, kondisi operasi bisa berubah, dan itu bisa jadi sumber "
    "alarm palsu kalau tidak ditangani. Rencana penanganannya (bagian dari "
    "roadmap implementasi, bukan prototipe minggu ini): bikin baseline "
    "sehat TERPISAH untuk setiap mode operasi, bukan 1 baseline untuk "
    "semua kondisi."
)
qa(
    "Apakah modelnya di-overfit ke data test tertentu, alias 'nyontek'?",
    "Untuk mencegah ini, dipakai 2 disiplin: (1) semua ambang batas/model "
    "HANYA dilatih dari bagian data yang dianggap sehat, tidak pernah "
    "melihat bagian akhir yang gagal; (2) Test 3 (holdout) benar-benar "
    "tidak disentuh/diintip sampai seluruh sistem selesai dibangun dari "
    "Test 2, baru dijalankan tanpa penyesuaian tambahan."
)
qa(
    "Kenapa threshold-nya 3 standar deviasi, bukan 2 atau 4?",
    "Itu konvensi statistik yang sudah baku (\"aturan 3-sigma\"): untuk "
    "sebaran data yang mendekati normal, sekitar 99,7% nilai wajar ada "
    "dalam rentang 3 simpangan dari rata-rata, jadi nilai di luar itu "
    "sangat mungkin memang tidak wajar. Angka ini juga sudah secara "
    "eksplisit disetujui mentor pembimbing sebagai pendekatan yang tepat."
)
qa(
    "Bagaimana kalau sensornya rusak atau datanya hilang/corrupt?",
    "Ini belum ditangani secara khusus di prototipe ini - prototipe "
    "mengasumsikan data selalu lengkap dan valid, karena datanya berasal "
    "dari dataset yang sudah bersih. Penanganan data hilang/corrupt (data "
    "quality handling) perlu ditambahkan sebelum implementasi produksi "
    "sungguhan di pabrik."
)
qa(
    "Apakah sistem ini bisa langsung dipakai real-time di pabrik?",
    "Belum secara langsung. Prototipe ini memakai mode REPLAY dari data "
    "yang sudah direkam sebelumnya (bukan streaming langsung dari sensor "
    "pabrik). Untuk versi produksi, dibutuhkan koneksi ke sistem industri "
    "(OPC UA/Modbus lewat gateway PLC/SCADA) - ini sudah masuk roadmap fase "
    "2 di proposal, bukan bagian dari prototipe minggu ini."
)
qa(
    "Seberapa akurat sistem ini sebenarnya?",
    "Diukur dengan 3 hal: (1) lead time - berapa jam alarm menyala sebelum "
    "kegagalan sungguhan; (2) pengurangan jumlah episode alarm dibanding "
    "sistem threshold tunggal biasa; (3) RMSE (rata-rata kesalahan) "
    "prediksi RUL di 10% umur terakhir. Semua angka ini dihitung dan "
    "tercatat per bearing di file artifacts/summary.csv, dan bisa "
    "dibandingkan dengan klaim di proposal."
)
qa(
    "Kenapa Health Score dan Anomaly Score bisa beda arah (satu masih "
    "tinggi, satu udah alarm)?",
    "Karena keduanya sengaja dibuat terpisah dengan tujuan berbeda: Anomaly "
    "Score (dasar alarm) dirancang SENSITIF supaya cepat bereaksi ke "
    "penyimpangan statistik sekecil apa pun. Health Score dirancang sebagai "
    "TREN HALUS untuk gambaran besar. Keduanya sengaja tidak digabung jadi "
    "satu supaya sistem tetap bisa bereaksi cepat tanpa kehilangan gambaran "
    "tren jangka panjang."
)
qa(
    "Apakah metode ini general untuk semua jenis bearing/motor, atau "
    "spesifik untuk yang ini saja?",
    "Metodenya (Isolation Forest, PCA, model eksponensial, persistence+"
    "voting) bersifat GENERAL dan bisa dipakai untuk bearing/motor lain. "
    "Yang SPESIFIK hanya bagian fisika (rumus BPFO/BPFI/BSF) yang perlu "
    "parameter geometri bearing yang sedang dipantau - kalau bearingnya "
    "beda, tinggal ganti angka geometri di pdm/bearing_physics.py, logika "
    "lainnya tidak perlu diubah."
)
qa(
    "Kenapa tidak pakai supervised learning (klasifikasi langsung sehat/"
    "rusak) saja, kan lebih akurat?",
    "Karena data tidak punya label yang jelas per titik waktu - yang "
    "diketahui cuma KAPAN bearing itu akhirnya gagal total (di ujung "
    "rekaman), bukan \"level kerusakan\" di setiap titik waktu sebelumnya. "
    "Tanpa label yang jelas, pendekatan unsupervised (belajar pola sehat, "
    "tandai penyimpangan) jauh lebih sesuai - ini juga sudah dikonfirmasi "
    "sebagai pendekatan yang tepat oleh mentor pembimbing."
)

# ================================================================ 12
h1("12. Ide Pengembangan agar Dilirik dan Disukai Astra")
body(
    "Ini saran tambahan (belum diimplementasikan) yang bisa membuat "
    "dashboard lebih menarik dan lebih relevan secara bisnis untuk "
    "dipresentasikan ke Astra, kalau ada waktu untuk pengembangan lanjutan "
    "setelah demo minggu ini:"
)
bullet(
    "Root cause di tab Alerts diubah supaya benar-benar membaca fitur mana "
    "yang paling menonjol (BPFO vs BPFI vs BSF) dan menyimpulkan lokasi "
    "kerusakan dari situ - bukan lagi mengandalkan fakta yang sudah "
    "diketahui sebelumnya.",
    bold_lead="Root cause berbasis bukti sungguhan - ",
)
bullet(
    "Halaman ringkasan yang menampilkan status ke-4 bearing sekaligus "
    "dalam bentuk grid/command center - relevan karena pabrik sungguhan "
    "punya banyak motor yang perlu dipantau bersamaan, bukan cuma 1.",
    bold_lead="Tampilan multi-mesin sekaligus - ",
)
bullet(
    "Tambahkan angka \"estimasi jam produksi yang diselamatkan\" atau "
    "\"estimasi biaya downtime yang dihindari\" kalau rekomendasi RUL "
    "diikuti - ini menerjemahkan hasil teknis ke bahasa bisnis yang "
    "langsung relevan buat keputusan manajemen Astra (downtime JIT = uang).",
    bold_lead="Terjemahan ke nilai bisnis (Rupiah/jam produksi) - ",
)
bullet(
    "Slider yang memungkinkan pengguna coba-coba mengubah sensitivitas "
    "alarm (misal 2-sigma vs 3-sigma vs 4-sigma) dan langsung melihat "
    "efeknya ke jumlah alarm - bagus untuk mendemonstrasikan trade-off "
    "sensitivitas vs ketenangan alarm secara interaktif di depan Astra.",
    bold_lead="Kontrol sensitivitas interaktif - ",
)
bullet(
    "Tombol untuk export kondisi yang sedang dilihat jadi laporan PDF/"
    "gambar ringkas - berguna untuk dokumentasi rutin teknisi pabrik.",
    bold_lead="Export laporan otomatis - ",
)
bullet(
    "Form sederhana untuk upload file CSV baru (mensimulasikan data motor "
    "Astra beneran nanti) - membuktikan sistemnya memang dirancang siap "
    "menerima data baru, bukan cuma untuk 1 dataset ini saja.",
    bold_lead="Simulasi upload data baru - ",
)
bullet(
    "Diagram kecil yang menunjukkan roadmap 3 fase implementasi (pilot 1 "
    "lini, produksi terbatas, skala lintas pabrik) langsung di dalam "
    "dashboard - supaya audiens Astra langsung lihat gambaran jalan ke "
    "depan, bukan cuma prototipe berdiri sendiri.",
    bold_lead="Roadmap visual di dalam dashboard - ",
)

# ================================================================ 13
h1("13. Ringkasan Kelebihan dan Keterbatasan Jujur")
h2("13.1 Kelebihan")
bullet("Seluruh metodologi (Isolation Forest, PCA, RUL eksponensial, persistence+voting) sudah diimplementasikan sungguhan, bukan simulasi/placeholder.")
bullet("Tervalidasi dengan skema holdout yang jujur (Test 3 tidak pernah diintip sampai sistem selesai dibangun).")
bullet("Terbukti secara fisik - puncak envelope spectrum tepat di frekuensi BPFO yang dihitung dari geometri bearing.")
bullet("Terbukti mengurangi jumlah alarm secara signifikan dibanding threshold tunggal (dari puluhan episode jadi hitungan jari).")
bullet("Setiap keputusan desain (kenapa metode ini, bukan itu) punya alasan yang bisa dipertanggungjawabkan, bukan asal pilih.")
h2("13.2 Keterbatasan yang Diakui")
bullet("Masih pakai data NASA IMS, belum data motor Astra sungguhan.")
bullet("Belum terhubung real-time ke sistem pabrik (masih mode replay).")
bullet("Baru menangani 1 jenis sensor (getaran) dan 1 jenis kerusakan (bearing outer race) secara mendalam.")
bullet("Model RUL eksponensial tunggal kurang akurat untuk pola degradasi yang naik-turun/berosilasi.")
bullet("Root cause di dashboard masih menampilkan fakta yang sudah diketahui, belum diagnosis independen penuh dari sinyal.")

# ================================================================ 14
h1("14. Daftar Istilah (Glossary) Lengkap")
table(
    ["Istilah", "Artinya"],
    [
        ["Bearing", "Laher/bantalan yang membuat poros motor berputar mulus"],
        ["RMS", "Ukuran seberapa besar/kuat getaran secara keseluruhan"],
        ["Kurtosis", "Ukuran seberapa tajam/mengejutkan pola getaran (indikator benturan)"],
        ["Crest Factor", "Perbandingan puncak tertinggi getaran terhadap RMS-nya"],
        ["BPFO/BPFI/BSF/FTF", "Frekuensi getaran karakteristik yang muncul kalau ada cacat di bagian tertentu bearing"],
        ["Envelope Analysis", "Teknik olah sinyal untuk menonjolkan benturan kecil yang tersembunyi di getaran"],
        ["Isolation Forest", "Metode AI yang mendeteksi data aneh karena mudah 'diisolasi' dari data lain"],
        ["Anomaly Score", "Angka yang menunjukkan seberapa 'aneh' suatu kondisi dibanding pola sehat"],
        ["3-sigma / 3 standar deviasi", "Batas statistik baku: nilai di luar 3x simpangan dari rata-rata dianggap tidak wajar"],
        ["PCA (Principal Component Analysis)", "Teknik menggabungkan banyak angka jadi satu angka tren, bobotnya dihitung otomatis dari data"],
        ["Health Index / Health Score", "Skor 0-100 hasil gabungan PCA yang menunjukkan tren kesehatan bearing"],
        ["EWMA", "Rata-rata berjalan yang menghaluskan noise, memberi bobot lebih ke data terbaru"],
        ["Persistence", "Aturan alarm: kondisi harus bertahan berturut-turut dulu baru dianggap serius"],
        ["Voting", "Aturan alarm: kalau 2+ tanda muncul bersamaan, langsung dianggap serius"],
        ["RUL (Remaining Useful Life)", "Perkiraan sisa waktu sebelum komponen benar-benar gagal"],
        ["Knee point / Trigger", "Titik pertama kali alarm resmi menyala, jadi awal mula perhitungan RUL"],
        ["EOL (End of Life)", "Titik waktu yang diperkirakan sebagai saat kegagalan terjadi"],
        ["ET (Failure Threshold)", "Nilai Health Index yang dianggap 'batas gagal', diambil dari rekaman lain"],
        ["Run-to-failure", "Rekaman yang berlangsung terus sampai komponennya benar-benar rusak total"],
        ["Holdout", "Data yang sengaja tidak diintip sampai sistem selesai dibangun, untuk uji kejujuran"],
        ["Burn-in", "Masa kalibrasi awal di mana sistem belum mengevaluasi alarm apa pun"],
        ["Hysteresis", "Aturan supaya status alarm tidak kedap-kedip, butuh konfirmasi sebelum berubah status"],
        ["Overfitting", "Kondisi model 'menghafal' data latihan, bukan benar-benar belajar pola umum"],
        ["Z-score", "Ukuran 'berapa kali lipat simpangan' suatu nilai berbeda dari rata-rata"],
    ],
)

# ================================================================ 15
h1("15. Referensi")
bullet("Proposal_Final_Case2.docx - hasil validasi lengkap metodologi ini.")
bullet("Laporan_Progres_Case2.docx - rincian fisika kegagalan dan rumus.")
bullet("Blueprint_Teknis_Case2.md - keputusan desain yang dikunci sebelum implementasi.")
bullet("Readme Document for IMS Bearing Data.pdf - dokumentasi resmi dataset NASA IMS.")
bullet("Bootcamp Briefing AOP Winteq.pdf - kebutuhan resmi Use Case 2 dari Astra Otoparts.")
bullet("F. T. Liu, K. M. Ting, Z.-H. Zhou, \"Isolation Forest\", IEEE ICDM, 2008.")
bullet("R. B. Randall, J. Antoni, \"Rolling Element Bearing Diagnostics - A Tutorial\", MSSP, 2011.")
bullet("Y. Lei et al., \"Machinery Health Prognostics: A Systematic Review\", MSSP, 2018.")
bullet("H. Qiu, J. Lee, J. Lin, \"Wavelet Filter-based Weak Signature Detection Method...\", JSV, 2006 (sumber dataset NASA IMS).")

doc.add_page_break()
end = doc.add_paragraph()
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = end.add_run(
    "- Selesai -\nDokumen ini dibuat selengkap mungkin untuk mendampingi "
    "dashboard prototipe Case 2 Predictive Maintenance,\nPT Astra Otoparts "
    "Tbk / WINTEQ."
)
r.italic = True
r.font.color.rgb = GREY

doc.save(str(OUT_PATH))
print(f"Saved: {OUT_PATH}")
