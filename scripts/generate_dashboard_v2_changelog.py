"""Generates docs/Penjelasan_Update_Dashboard_FleetOverview_Case2.docx - a
complete, plain-language explanation of every change made to the dashboard
in the "Fleet Overview + alarm + synthetic sensors + visual redesign" round
(2026-07-17 session), for Jingga to read/forward to her supervisor.

Run:  python -m scripts.generate_dashboard_v2_changelog   (from src/)
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm

OUT_PATH = Path(__file__).resolve().parents[1] / "docs" / "Penjelasan_Update_Dashboard_FleetOverview_Case2.docx"

NAVY = RGBColor(0x0F, 0x17, 0x2A)
BLUE = RGBColor(0x03, 0x69, 0xA1)
GREY = RGBColor(0x47, 0x55, 0x69)
GREEN = RGBColor(0x16, 0xA3, 0x4A)
RED = RGBColor(0xDC, 0x26, 0x26)
ORANGE = RGBColor(0xD9, 0x77, 0x06)

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def set_margins():
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)


set_margins()


def h1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = NAVY
    return p


def h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = BLUE
    return p


def h3(text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = GREY
    return p


def body(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(8)
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def numbered(text, bold_lead=None):
    p = doc.add_paragraph(style="List Number")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def note_box(text, color=BLUE, label="CATATAN"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(f"{label}: " + text)
    run.italic = True
    run.font.color.rgb = color
    return p


def simple_table(headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, htext in enumerate(headers):
        hdr_cells[i].text = htext
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


def status_table(headers, rows):
    """Same as simple_table but colors the first cell of each row by an
    implicit OK/PARTIAL/PENDING marker embedded as the first char of val."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, htext in enumerate(headers):
        hdr_cells[i].text = htext
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    color_map = {"[OK]": GREEN, "[SEBAGIAN]": ORANGE, "[BELUM]": RED}
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            text = str(val)
            color = None
            for marker, c in color_map.items():
                if text.startswith(marker):
                    color = c
                    text = text[len(marker):].strip()
                    break
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(text)
            if color:
                r.font.color.rgb = color
                r.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


# ============================================================ COVER
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("PENJELASAN LENGKAP")
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Update Dashboard: Fleet Overview, Alarm, Data Sensor Tambahan, dan Tampilan Baru")
run.font.size = Pt(15)
run.font.color.rgb = BLUE

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub2.add_run("Case 2 - Bootcamp AOP Winteq · PT Astra Otoparts Tbk")
run.font.size = Pt(12)
run.font.color.rgb = GREY

doc.add_paragraph()
note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = note.add_run(
    "Dokumen ini menjelaskan SEMUA perubahan yang baru ditambahkan ke dashboard - "
    "bukan dari nol, tapi tambahan/perbaikan dari dashboard versi sebelumnya. Ditulis "
    "sesederhana mungkin, termasuk penjelasan istilah-istilah teknis, supaya bisa "
    "dipahami dan dijelaskan ulang ke siapa saja tanpa perlu latar belakang teknik."
)
run.italic = True
run.font.color.rgb = GREY

doc.add_page_break()

# ============================================================ DAFTAR ISI
h1("Daftar Isi")
toc_items = [
    "1. Ringkasan Singkat (Baca Ini Dulu Kalau Buru-Buru)",
    "2. Kenapa Perubahan Ini Dibuat?",
    "3. Perubahan #1 - Halaman Baru: Fleet Overview",
    "4. Perubahan #2 - Alarm CRITICAL yang Mencolok (Pop-up + Suara)",
    "5. Perubahan #3 - Data Tambahan: Temperature dan Current (INI YANG PALING PENTING DIPAHAMI)",
    "6. Perubahan #4 - Filter Berdasarkan Parameter",
    "7. Perubahan #5 - Status Normal/Elevated per Parameter (Biar Tidak Cuma Pajangan)",
    "8. Perubahan #6 - Jendela \"Lihat Cepat\" (Quick View)",
    "9. Perubahan #7 - Tampilan Baru (Logo, Warna, Pattern, Font)",
    "10. Bug yang Ditemukan dan Diperbaiki",
    "11. Tabel Perbandingan: Sebelum vs Sesudah",
    "12. Apakah Sudah Sesuai dengan Permintaan Astra?",
    "13. Apakah Sudah Sesuai dengan Feedback Dosen?",
    "14. Hal-Hal yang Harus Jujur Disampaikan (Kalau Ditanya)",
    "15. Daftar Istilah (Glossary)",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================ 1. RINGKASAN
h1("1. Ringkasan Singkat (Baca Ini Dulu Kalau Buru-Buru)")
body(
    "Dashboard yang tadinya cuma bisa menampilkan SATU mesin/bearing dalam satu "
    "waktu (harus pilih dulu di sidebar), sekarang ditambah dengan halaman baru "
    "yang menampilkan SEMUA mesin sekaligus dalam bentuk kartu-kartu, supaya "
    "manajer atau pengawas bisa langsung tahu mesin mana yang bermasalah tanpa "
    "harus cek satu-satu."
)
bullet("Halaman baru \"Fleet Overview\" - lihat semua 8 bearing sekaligus dalam bentuk kartu, seperti command center.", bold_lead="Yang baru: ")
bullet("Kalau ada mesin CRITICAL, muncul banner merah besar berkedip + suara alarm yang terus berbunyi sampai di-\"Acknowledge\" (ditekan tombolnya).", bold_lead="Yang baru: ")
bullet("Ditambahkan data suhu (temperature) dan arus listrik (current) sebagai pelengkap data getaran (vibration) yang sudah ada - tapi ini data BUATAN/SIMULASI, bukan data sensor asli (dijelaskan detail di Bagian 5).", bold_lead="Yang baru: ")
bullet("Supaya data suhu dan arus listrik itu tidak cuma jadi pajangan, sekarang tiap parameter punya label \"Normal\"/\"Elevated\" sendiri, plus 1 kalimat otomatis yang membandingkan ketiganya (Bagian 7).", bold_lead="Yang baru: ")
bullet("Tampilan visual dibuat lebih \"mewah\": logo Astra Otoparts asli, warna dan bayangan yang lebih hidup, background yang bergerak pelan, angka-angka pakai font khusus yang kesannya presisi/teknis.", bold_lead="Yang baru: ")
bullet("Ditemukan dan diperbaiki 3 bug asli selama proses testing (bukan cuma dugaan - semua sudah dicoba langsung dan terbukti benar): tombol \"Acknowledge\" yang tidak bisa diklik saat mode Play jalan, error warna pada grafik, dan error nge-save posisi halaman.", bold_lead="Yang diperbaiki: ")
note_box(
    "Metodologi inti (cara dashboard menghitung Health Score, kapan alarm menyala, "
    "dan estimasi RUL/Remaining Useful Life) SAMA SEKALI TIDAK DIUBAH. Semua "
    "perubahan di atas sifatnya nambah, bukan mengganti otak utamanya. Jadi angka-"
    "angka yang sudah pernah dipresentasikan (health score 83.1 di checkpoint demo, "
    "dan lain-lain) masih tetap sama persis.",
    color=GREEN, label="PENTING",
)

doc.add_page_break()

# ============================================================ 2. KENAPA
h1("2. Kenapa Perubahan Ini Dibuat?")
body(
    "Semua perubahan ini berawal dari feedback dosen setelah presentasi Week 2. "
    "Ada 2 poin utama yang diminta, plus beberapa poin tambahan:"
)
h3("2.1 Poin Utama dari Dosen")
numbered(
    "kalau ada 300 mesin, seorang manajer tidak mungkin cek satu-satu dari sidebar. "
    "Harus ada 1 halaman yang langsung menunjukkan mesin mana yang mau rusak, "
    "dalam format sederhana seperti \"Mesin X - Bearing Y → critical dalam 2 jam\".",
    bold_lead="Tambah menu \"Dashboard\"/high-level view - ",
)
numbered(
    "untuk kondisi CRITICAL, minta ada pop-up besar yang muncul berulang sampai "
    "operator benar-benar menindaklanjuti (klik acknowledge), boleh ditambah suara "
    "alarm juga.",
    bold_lead="Notifikasi/alarm lebih mencolok (\"annoying\") - ",
)
h3("2.2 Poin Tambahan")
bullet(
    "boleh dipertimbangkan bikin data sintetis (buatan/simulasi) untuk parameter "
    "lain selain vibrasi, karena NASA IMS Bearing Dataset yang dipakai memang "
    "cuma punya data getaran.",
    bold_lead="Data tambahan (temperature, dll) - ",
)
bullet(
    "kalau ada parameter tambahan, sebaiknya ada filter di dashboard berdasarkan "
    "parameter itu (vibration-based, temperature-based, dst).",
    bold_lead="Filter berdasarkan parameter - ",
)
bullet(
    "user juga minta ada cara untuk \"intip\" data satu mesin sebelum benar-benar "
    "masuk ke halaman detail lengkapnya, dan minta tampilan dashboard dibuat lebih "
    "profesional/menarik secara visual.",
    bold_lead="Permintaan tambahan dari user (Jingga) - ",
)
body(
    "Semua poin di atas sudah dikerjakan. Detail masing-masing ada di bagian-bagian "
    "berikutnya."
)

doc.add_page_break()

# ============================================================ 3. FLEET OVERVIEW
h1("3. Perubahan #1 - Halaman Baru: Fleet Overview")
body(
    "Ini halaman BARU yang sekarang jadi halaman pertama yang muncul saat dashboard "
    "dibuka (sebelumnya, yang pertama muncul langsung detail 1 mesin). Analoginya "
    "seperti ruang kontrol di pabrik yang menampilkan status SEMUA mesin di satu "
    "layar, baru kalau ada yang aneh, operator baru \"zoom in\" ke mesin itu."
)
h2("3.1 Apa saja isinya?")
bullet(
    "1 slider yang menggerakkan SEMUA 8 bearing bersamaan ke titik waktu yang "
    "relatif sama (misalnya 78% berarti \"78% dari total waktu berjalannya masing-"
    "masing mesin\"). Karena Test 2 berjalan 163.8 jam dan Test 3 berjalan 1073.3 "
    "jam (jauh lebih lama), pakai persentase supaya kedua kelompok mesin bisa "
    "\"disamakan waktunya\" walau panjang datanya beda jauh.",
    bold_lead="Fleet Time (penggeser waktu bersama) - ",
)
bullet(
    "dropdown untuk memilih mau lihat data \"All parameters\", \"Vibration\", "
    "\"Temperature\", atau \"Current\" saja pada tiap kartu.",
    bold_lead="Parameter focus - ",
)
bullet(
    "4 kotak angka besar: total mesin dipantau, jumlah yang Critical, jumlah "
    "Warning, jumlah Normal - supaya dalam sekali lihat, langsung tahu skala "
    "masalahnya.",
    bold_lead="KPI strip (ringkasan angka) - ",
)
bullet(
    "daftar mesin yang statusnya bukan Normal, diurutkan dari yang paling parah, "
    "dengan format persis seperti contoh dosen: \"Test 2 - Bearing 1 → Critical, "
    "failure expected in 1d 5h\". Kalau semua mesin normal, muncul pesan aman "
    "berwarna hijau.",
    bold_lead="\"Needs attention\" (daftar prioritas) - ",
)
bullet(
    "kartu untuk SEMUA 8 bearing (dikelompokkan per Test), tiap kartu menunjukkan "
    "status (Normal/Warning/Critical dengan warna), health score, estimasi RUL, "
    "dan data vibrasi/temperature/current sesuai filter yang dipilih. Ada juga "
    "centang \"Show only Warning/Critical machines\" untuk menyembunyikan mesin "
    "yang sehat kalau mau fokus ke yang bermasalah saja.",
    bold_lead="\"All machines\" (grid semua mesin) - ",
)
note_box(
    "Kenapa nilai default slider-nya 78%? Karena di titik itu, hasilnya persis "
    "menunjukkan Test 2 Bearing 1 dalam kondisi Critical (ini bearing yang sama "
    "yang dipakai di checkpoint demo lama kalian) sementara mesin lain masih "
    "Normal/Warning - jadi begitu dashboard dibuka, langsung ada contoh nyata "
    "yang bisa didemokan tanpa perlu geser-geser slider dulu."
)

doc.add_page_break()

# ============================================================ 4. ALARM CRITICAL
h1("4. Perubahan #2 - Alarm CRITICAL yang Mencolok (Pop-up + Suara)")
body(
    "ini persis permintaan dosen soal notifikasi yang \"annoying\" - dibuat supaya "
    "operator TIDAK MUNGKIN tidak sadar kalau ada mesin yang statusnya sudah "
    "CRITICAL."
)
h2("4.1 Bagaimana cara kerjanya")
numbered(
    "Kalau ada mesin berstatus CRITICAL dan belum di-\"acknowledge\" (belum "
    "ditanggapi), sebuah banner merah besar muncul di paling atas layar, "
    "menutupi lebar layar penuh."
)
numbered(
    "Banner ini BERKEDIP (warnanya berubah dari merah ke merah tua terus-menerus) "
    "supaya menarik perhatian mata."
)
numbered(
    "Bersamaan dengan itu, ada suara alarm (bunyi sirine 2 nada bergantian) yang "
    "TERUS BERULANG selama banner masih muncul."
)
numbered(
    "Di bawah banner ada tombol \"Acknowledge & continue monitoring\". Begitu "
    "diklik, banner dan suaranya hilang, dan status itu dianggap sudah "
    "ditanggapi."
)
numbered(
    "Kalau nanti kondisinya membaik lalu memburuk lagi (atau ada mesin LAIN yang "
    "jadi Critical), banner akan muncul lagi dari awal - jadi tidak mungkin ada "
    "alarm baru yang terlewat hanya karena alarm sebelumnya pernah di-acknowledge."
)
body(
    "Suara alarmnya dibuat sendiri langsung lewat kode (bukan file audio dari "
    "internet) - jadi tidak ada file tambahan yang perlu di-download, dan tidak "
    "ada risiko file hilang/rusak."
)
note_box(
    "Browser modern (Chrome, Edge, dst) punya aturan keamanan: suara TIDAK BOLEH "
    "otomatis berbunyi sebelum ada klik apa pun dari user di halaman itu. Jadi "
    "kalau dashboard baru dibuka dan langsung ada alarm tapi belum terdengar "
    "suaranya, klik dulu di mana saja pada halaman - baru alarm berikutnya akan "
    "kedengaran. Ini bukan bug, ini memang aturan dari browser, di luar kendali "
    "aplikasi manapun."
)

doc.add_page_break()

# ============================================================ 5. SYNTHETIC DATA
h1("5. Perubahan #3 - Data Tambahan: Temperature dan Current")
body(
    "INI BAGIAN YANG PALING PENTING UNTUK DIPAHAMI, karena paling mungkin "
    "ditanya dosen atau reviewer Astra. Dibaca pelan-pelan ya."
)
h2("5.1 Apa itu dan kenapa ada?")
body(
    "Dataset NASA IMS Bearing yang dipakai sebagai pengganti data Astra (karena "
    "Astra belum kasih data sensor asli) itu CUMA punya satu jenis data: getaran "
    "(vibration). Padahal permintaan Astra di soal Case 2 minta pemantauan "
    "MULTI-PARAMETER: suhu (temperature), getaran (vibration), arus listrik "
    "(current), tekanan (pressure), dan lain-lain."
)
body(
    "Supaya dashboard bisa menunjukkan kesiapan untuk multi-parameter (sesuai "
    "permintaan Astra), ditambahkan dua kolom data baru: temperature_c (suhu, "
    "dalam derajat Celsius) dan current_a (arus listrik, dalam Ampere)."
)
note_box(
    "Data ini BUATAN/SIMULASI (kita sebut \"data sintetis\"), BUKAN hasil "
    "pengukuran sensor asli - karena memang tidak ada sensor suhu/arus yang "
    "sungguhan dipasang. Ini WAJIB disampaikan dengan jujur kalau ditanya. "
    "Jangan pernah bilang ini data sensor asli.",
    color=RED, label="PENTING - JANGAN LUPA",
)
h2("5.2 Bagaimana cara data ini dibuat? (supaya tidak asal-asalan)")
body(
    "Data ini TIDAK dibuat acak/random begitu saja. Cara buatnya mengikuti logika "
    "yang masuk akal secara fisika, yaitu:"
)
bullet(
    "dashboard sudah lebih dulu menghitung skor \"seberapa rusak\" bearing itu "
    "dari data getaran asli (nilai ini namanya hi_norm, bagian dari perhitungan "
    "Health Index yang sudah ada sejak awal dan TIDAK diubah).",
    bold_lead="Langkah 1 - ",
)
bullet(
    "dari skor kerusakan itu, dihitung nilai suhu dan arus listrik yang IKUT NAIK "
    "kalau bearing-nya makin rusak. Ini masuk akal secara fisika: bearing yang "
    "rusak menghasilkan gesekan lebih besar → gesekan menghasilkan panas "
    "(suhu naik) dan motor butuh tenaga listrik lebih besar untuk tetap berputar "
    "(arus naik).",
    bold_lead="Langkah 2 - ",
)
bullet(
    "suhu naiknya lebih LAMBAT dibanding getaran (karena panas butuh waktu untuk "
    "menumpuk, tidak langsung naik seketika), dan arus listrik naiknya lebih "
    "lambat dan lebih kecil/tidak stabil lagi dibanding suhu. Ini supaya tidak "
    "terlihat seperti 3 angka yang gerakannya identik - di dunia nyata, tiap "
    "sensor punya karakter responsnya sendiri-sendiri.",
    bold_lead="Langkah 3 (beda kecepatan tiap sensor) - ",
)
bullet(
    "ditambahkan sedikit noise/gangguan kecil supaya angkanya tidak terlihat "
    "terlalu rapi/matematis - persis seperti sensor sungguhan yang selalu ada "
    "sedikit \"getaran\" di angkanya.",
    bold_lead="Langkah 4 - ",
)
simple_table(
    ["Parameter", "Suhu normal (sehat)", "Kenaikan maksimum", "Kecepatan respons"],
    [
        ["Temperature (suhu)", "≈ 38°C", "sampai ≈ 70°C", "Sedang (ada jeda/lag)"],
        ["Current (arus listrik)", "≈ 4.1 A", "sampai ≈ 7.1 A", "Paling lambat & paling tidak stabil"],
    ],
)
h2("5.3 Ini bagian yang paling sering bikin bingung: apakah data ini ikut menentukan Health Score?")
body("JAWABANNYA: TIDAK.", bold=True)
body(
    "Arahnya SATU ARAH SAJA, seperti ini:"
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Data Getaran (asli)  →  Health Score, Alarm, RUL (sudah final)  →  Temperature & Current (buatan)")
r.bold = True
r.font.color.rgb = BLUE
p.paragraph_format.space_after = Pt(10)
body(
    "Artinya: Health Score, status alarm (Normal/Warning/Critical), dan estimasi "
    "RUL itu SEMUA masih 100% dihitung dari data getaran saja, TIDAK BERUBAH "
    "SAMA SEKALI dari sebelumnya. Data suhu dan arus listrik itu dibuat BELAKANGAN, "
    "dari hasil yang sudah jadi - jadi sifatnya cuma \"ikut nempel\" untuk "
    "ditampilkan, bukan bahan hitungan."
)
body(
    "Analoginya seperti spreadsheet Excel: kolom A isinya angka asli, kolom B "
    "isinya rumus \"=A x 2\". Kolom B ikut berubah kalau kolom A berubah, tapi "
    "kolom A tidak butuh kolom B untuk dihitung. Nah, Health Score itu kolom A "
    "(sudah final dan independen), sedangkan suhu/arus listrik itu kolom B "
    "(mengikuti, bukan menentukan)."
)
note_box(
    "Kenapa sengaja dibuat begini (bukan saling mempengaruhi)? Supaya SEMUA "
    "angka yang sudah pernah diuji dan dipresentasikan sebelumnya (lead time, "
    "RMSE RUL, jumlah episode alarm, dll di artifacts/summary.csv) tetap 100% "
    "sama, tidak ikut berubah/rusak gara-gara data karangan. Kalau data buatan "
    "ini malah dimasukkan ke perhitungan Health Score, semua hasil yang sudah "
    "divalidasi bisa berubah dan jadi tidak bisa dipertanggungjawabkan lagi.",
    color=GREEN,
)
h2("5.4 Di mana bisa dilihat buktinya?")
bullet(
    "berisi kode lengkap dengan penjelasan cara kerja di setiap barisnya.",
    bold_lead="pdm/synthetic_sensors.py - ",
)
bullet(
    "kolom temperature_c dan current_a bisa dibuka langsung, satu baris data per "
    "10 menit, sama seperti kolom-kolom lain yang sudah ada.",
    bold_lead="artifacts/test2_bearing1.csv (dan file lain sejenis) - ",
)
bullet(
    "kalimat \"Temperature and current are synthetic...\" muncul di halaman "
    "Fleet Overview dan di tab \"Sensors\" pada Machine Detail - supaya siapa "
    "pun yang buka dashboard langsung tahu ini bukan data sensor asli tanpa "
    "perlu ditanya dulu.",
    bold_lead="Tulisan keterangan di dashboard - ",
)

doc.add_page_break()

# ============================================================ 6. FILTER
h1("6. Perubahan #4 - Filter Berdasarkan Parameter")
body(
    "Ini menjawab permintaan tambahan dosen: \"kalau ada parameter tambahan, "
    "tambahin juga filter di dashboard berdasarkan parameter tersebut.\""
)
bullet(
    "dropdown \"Parameter focus\" (All parameters / Vibration / Temperature / "
    "Current) - mengganti data tambahan apa yang ditampilkan di tiap kartu mesin.",
    bold_lead="Di halaman Fleet Overview - ",
)
bullet(
    "tab baru bernama \"Sensors\" berisi 3 grafik (Vibration, Temperature, "
    "Current) masing-masing dengan garis putus-putus merah yang menunjukkan "
    "\"batas normal\" (baseline sehat). Ada kotak centang untuk memilih grafik "
    "mana saja yang mau ditampilkan/disembunyikan.",
    bold_lead="Di halaman Machine Detail - ",
)
note_box(
    "Status alarm (Normal/Warning/Critical) TETAP hanya ditentukan oleh data "
    "getaran, sesuai penjelasan di Bagian 5.3. Filter ini murni untuk keperluan "
    "TAMPILAN saja, tidak mengubah logika deteksi."
)

doc.add_page_break()

# ============================================================ 7. STATUS PER PARAMETER
h1("7. Perubahan #5 - Status Normal/Elevated per Parameter (Biar Tidak Cuma Pajangan)")
body(
    "Perubahan ini muncul karena pertanyaan yang sangat tepat: \"kalau temperature "
    "dan current cuma ditampilkan tapi tidak menentukan apa-apa, apa gunanya? "
    "Bukannya jadi pajangan doang?\" Jawabannya: SEBELUM perubahan ini, iya, betul, "
    "temperature dan current itu memang cuma dipajang. Bagian ini menambahkan "
    "\"fungsi\" ke keduanya - TANPA mengorbankan kejujuran bahwa data itu sintetis, "
    "dan TANPA membuatnya ikut menentukan alarm resmi (supaya angka yang sudah "
    "divalidasi tetap aman, sesuai penjelasan Bagian 5.3)."
)
h2("7.1 Apa yang ditambahkan?")
bullet(
    "di tab \"Sensors\" (Machine Detail), tiap grafik (Vibration, Temperature, "
    "Current) sekarang punya label kecil di pojok kanan atas: \"NORMAL\" (hijau) "
    "atau \"ELEVATED\" (oranye, artinya \"lagi naik/di atas kebiasaan\"). Label "
    "ini dihitung dengan membandingkan nilai parameter itu SAAT INI dengan batas "
    "3-sigma (baseline sehat) milik parameter itu sendiri - jadi tiap parameter "
    "dinilai berdasarkan kebiasaannya masing-masing, bukan disamaratakan.",
    bold_lead="Badge Normal/Elevated - ",
)
bullet(
    "di atas ketiga grafik, muncul 1 kalimat otomatis yang membandingkan ketiga "
    "status itu sekaligus dan menyimpulkan artinya. Kalimat ini berubah sendiri "
    "tergantung kombinasi status yang sedang terjadi - ada 3 kemungkinan pola "
    "yang bisa muncul (lihat tabel di bawah).",
    bold_lead="Kalimat ringkasan (\"agreement summary\") - ",
)
h2("7.2 Contoh kalimat yang bisa muncul")
simple_table(
    ["Kondisi", "Kalimat yang muncul", "Artinya"],
    [
        [
            "Semua parameter normal",
            "\"All three parameters are within their normal range - consistent with a healthy bearing.\"",
            "Tidak ada tanda-tanda masalah sama sekali di ketiga data.",
        ],
        [
            "Cuma vibrasi yang naik",
            "\"Only vibration is elevated so far - temperature and current have not followed yet...\"",
            "Wajar untuk kerusakan tahap awal - getaran memang biasanya yang PERTAMA bereaksi, sebelum panas/arus listrik ikut naik.",
        ],
        [
            "Vibrasi + temperature + current semua naik bareng",
            "\"Vibration AND Temperature and Current are elevated together right now - a stronger, corroborating signal...\"",
            "Sinyal yang lebih 'meyakinkan' karena bukan cuma 1 data yang bergerak, tapi konsisten di beberapa data sekaligus.",
        ],
    ],
)
h2("7.3 Kenapa ini tetap aman, tidak mengulang masalah data karangan yang dijelaskan di Bagian 5?")
body(
    "Karena badge dan kalimat ini SIFATNYA CUMA MEMBACA ULANG data yang sudah "
    "ada dan menampilkannya dengan cara yang lebih mudah dimengerti - bukan "
    "menghitung sesuatu yang baru, dan sama sekali tidak menyentuh perhitungan "
    "Health Score/Alarm/RUL resmi. Ibaratnya seperti termometer badan yang "
    "layarnya menambahkan tulisan \"Demam\" kalau suhunya di atas 37.5°C - "
    "tulisan itu membantu orang awam membaca angka lebih cepat, tapi tidak "
    "mengubah suhu badan yang sebenarnya."
)
note_box(
    "Badge \"ELEVATED\" ini adalah pembacaan SEDERHANA (cuma bandingkan dengan "
    "batas sendiri-sendiri), BUKAN pengganti logika alarm resmi yang jauh lebih "
    "canggih (yang menggabungkan banyak fitur getaran sekaligus lewat aturan "
    "persistence + voting, dijelaskan di dokumen lain). Jadi wajar kalau "
    "sesekali badge ini menyala \"ELEVATED\" padahal status alarm resmi masih "
    "\"Warning\" (belum \"Critical\") - itu bukan bug, itu memang dua cara "
    "membaca data yang berbeda tingkat kerumitannya."
)

doc.add_page_break()

# ============================================================ 8. QUICK VIEW
h1("8. Perubahan #6 - Jendela \"Lihat Cepat\" (Quick View)")
body(
    "Ini jawaban dari permintaan: \"aku mau bisa lihat data dulu sebelum masuk "
    "ke halaman detail penuh, tanpa harus pindah halaman.\""
)
numbered("Di setiap kartu mesin (baik di daftar \"Needs attention\" maupun grid \"All machines\"), ada tombol \"Quick view\".")
numbered(
    "Begitu diklik, muncul jendela kecil MELAYANG di tengah layar (halaman Fleet "
    "Overview di belakangnya jadi agak redup, tapi tidak hilang) - isinya: nama "
    "mesin, status, health score, estimasi RUL, grafik mini tren health score, "
    "dan alasan alarm."
)
numbered(
    "Ada 2 tombol di jendela itu: \"Open full detail\" (baru pindah ke halaman "
    "Machine Detail lengkap) atau \"Close\" (tutup jendela, balik ke Fleet "
    "Overview persis seperti semula, posisi slider dan filter tidak berubah)."
)
body(
    "Istilah teknisnya, jendela melayang seperti ini disebut \"modal\" atau "
    "\"dialog\" - pola desain yang umum dipakai di aplikasi modern (misalnya "
    "pop-up konfirmasi saat mau hapus sesuatu di aplikasi HP)."
)

doc.add_page_break()

# ============================================================ 8. VISUAL
h1("9. Perubahan #7 - Tampilan Baru (Logo, Warna, Pattern, Font)")
body(
    "Bagian ini murni soal tampilan (bukan soal logika/perhitungan), supaya "
    "dashboard terlihat lebih profesional dan tidak polos."
)
h2("8.1 Logo Astra Otoparts")
body(
    "Logo asli PT Astra Otoparts Tbk (globe biru dengan garis merah, tulisan "
    "\"ASTRA Otoparts\") dipasang di dashboard, menggantikan tulisan biasa yang "
    "dipakai sebelumnya."
)
bullet("logo lengkap dengan tulisan \"ASTRA Otoparts\", karena latar belakangnya putih/terang sehingga logo aslinya (merah-hitam) kelihatan jelas.", bold_lead="Di sidebar (panel kiri) - ")
bullet("hanya ikon bulat globe-nya saja (tanpa tulisan, tanpa kotak putih di belakangnya), karena permintaan khusus supaya tidak terlalu ramai di bar judul yang warnanya gelap.", bold_lead="Di setiap judul halaman (topbar) - ")
note_box(
    "Sempat ada 2 kali revisi logo: pertama sempat pakai gambar buatan sendiri "
    "(karena awalnya belum ada file logo asli), lalu diganti logo Astra "
    "International (salah, itu logo perusahaan induk bukan Astra Otoparts), "
    "baru terakhir diganti logo Astra Otoparts yang benar sesuai file yang "
    "diberikan."
)
h2("8.2 Warna, bayangan, dan bentuk kartu")
bullet("kartu-kartu (KPI, machine card) sekarang punya efek bayangan berlapis dan sedikit \"naik\" saat mouse diarahkan ke situ (hover), bukan lagi kotak datar polos.", bold_lead="Shadow (bayangan) - ")
bullet("bar judul di atas tiap halaman sekarang pakai gradasi 3 warna (bukan cuma 2) plus tekstur titik-titik halus, supaya terlihat lebih \"berkelas\" seperti aplikasi enterprise.", bold_lead="Gradient (gradasi warna) - ")
bullet("label status \"Critical\" sekarang berkedip pelan (glow/pulse) supaya lebih menarik perhatian, tidak cuma diam.", bold_lead="Badge status - ")
bullet("angka-angka penting (health score, RUL, suhu, arus) sekarang pakai font \"JetBrains Mono\" - font khusus yang biasa dipakai di aplikasi teknis/coding, kesannya lebih presisi dan \"insinyur\", beda dari font tulisan biasa.", bold_lead="Font angka (monospace) - ")
bullet("grafik-grafik sekarang punya area terisi warna tipis di bawah garis (bukan cuma garis putih polos), dan garis bantu (grid) yang lebih halus.", bold_lead="Grafik - ")
h2("8.3 Background bergerak (pattern)")
body(
    "Latar belakang seluruh halaman sekarang punya pola garis-garis membentuk "
    "pola heksagonal/segitiga (seperti kisi-kisi teknis/blueprint), plus dua "
    "titik cahaya lembut yang bergeser pelan-pelan. Semuanya bergerak SANGAT "
    "PELAN dan LEMBUT (siklus sekitar 26 detik), supaya halaman terasa \"hidup\" "
    "tanpa mengganggu saat membaca data."
)
note_box(
    "Kalau di pengaturan komputer/browser ada opsi \"reduce motion\" atau "
    "\"kurangi animasi\" (biasanya untuk orang yang mudah pusing lihat animasi), "
    "background ini otomatis berhenti bergerak - ini standar aksesibilitas, "
    "bukan bug."
)

doc.add_page_break()

# ============================================================ 9. BUGS
h1("10. Bug yang Ditemukan dan Diperbaiki")
body(
    "Semua bug di bawah ini BENAR-BENAR ditemukan lewat percobaan langsung "
    "(dashboard dijalankan dan diklik-klik pakai program otomatis untuk "
    "memastikan semuanya benar-benar berfungsi, bukan cuma dugaan/asumsi), dan "
    "semuanya sudah dites ulang setelah diperbaiki untuk memastikan benar-benar "
    "beres."
)
h2("9.1 Tombol \"Acknowledge\" tidak bisa diklik saat mode Play jalan (BUG PALING PENTING)")
body(
    "Ini bug yang dilaporkan langsung dari pengalaman pakai dashboard: saat "
    "tombol Play ditekan dan animasi berjalan otomatis (auto-play), begitu "
    "statusnya sampai CRITICAL dan muncul banner alarm, tombol Acknowledge-nya "
    "TIDAK BISA diklik - harus refresh halaman untuk menghentikannya."
)
h3("Kenapa bisa terjadi?")
body(
    "Saat mode Play aktif, dashboard menyegarkan/menggambar ulang seluruh "
    "halaman secara otomatis berkali-kali per detik (untuk menggerakkan slider "
    "posisi waktu). Setiap kali halaman digambar ulang, tombol Acknowledge juga "
    "ikut \"dibongkar dan dipasang lagi\" dari awal. Klik dari user butuh waktu "
    "sepersekian detik untuk benar-benar terkirim dan diproses - tapi karena "
    "halaman keburu digambar ulang lagi sebelum klik itu selesai diproses, "
    "kliknya jadi seperti tidak pernah terjadi. Ibaratnya seperti mencoba naik "
    "eskalator yang bergerak terlalu cepat - keburu berganti sebelum kaki "
    "sempat menapak."
)
h3("Bagaimana cara memperbaikinya?")
body(
    "Sekarang, begitu status masuk CRITICAL dan banner-nya muncul (belum "
    "di-acknowledge), mode Play OTOMATIS BERHENTI SENDIRI. Jadi halaman tidak "
    "lagi menggambar ulang terus-menerus, dan tombol Acknowledge jadi stabil "
    "dan bisa diklik dengan normal. Ini juga sekaligus lebih masuk akal secara "
    "logika bisnis: sistem alarm sungguhan memang seharusnya \"berhenti dan "
    "menunggu\" sampai ditanggapi, bukan terus jalan seolah tidak terjadi apa-apa."
)
h3("Bagaimana cara memastikan sudah benar-benar beres?")
body(
    "Dites dengan program otomatis: geser slider ke posisi jauh sebelum alarm "
    "muncul, tekan Play, biarkan berjalan sampai masuk CRITICAL, lalu diperiksa "
    "apakah jam berjalannya benar-benar berhenti (dicek 2 kali dengan jeda "
    "2.5 detik - hasilnya sama persis, artinya sudah tidak jalan lagi), lalu "
    "dicoba klik tombol Acknowledge - berhasil, dan alarm-nya hilang."
)
h2("9.2 Error warna pada grafik area terisi")
body(
    "Saat menambahkan efek \"area terisi warna\" di bawah garis grafik, sempat "
    "muncul error karena format kode warna yang dipakai (kode warna 8 digit "
    "dengan tingkat transparansi) tidak dikenali oleh Plotly (library pembuat "
    "grafik yang dipakai dashboard). Sudah diperbaiki dengan mengganti ke format "
    "warna yang benar (format rgba)."
)
h2("9.3 Error saat menekan tombol \"View details\"/navigasi dari kartu")
body(
    "Sempat muncul error \"StreamlitAPIException\" saat tombol navigasi dari "
    "kartu mesin ke halaman detail ditekan, karena ada aturan di Streamlit "
    "(perangkat lunak yang dipakai bikin dashboard ini) bahwa suatu nilai "
    "pengaturan halaman tidak boleh diubah langsung setelah tombol pilihannya "
    "sudah ditampilkan di layar. Sudah diperbaiki dengan mengubah dulu nilainya "
    "SEBELUM tombol pilihan itu ditampilkan (pola yang sama seperti yang sudah "
    "dipakai di bagian lain dashboard untuk mengatur slider posisi)."
)

doc.add_page_break()

# ============================================================ 10. BEFORE/AFTER
h1("11. Tabel Perbandingan: Sebelum vs Sesudah")
simple_table(
    ["Aspek", "Sebelum", "Sesudah"],
    [
        ["Halaman utama", "Langsung ke detail 1 mesin, harus pilih dari sidebar", "Fleet Overview (semua mesin) jadi halaman pertama"],
        ["Lihat semua mesin sekaligus", "Tidak bisa, harus buka satu-satu", "Bisa, dalam bentuk grid kartu"],
        ["Alarm Critical", "Toast kecil sekali muncul (notifikasi tipis di pojok)", "Banner besar berkedip + suara alarm, wajib di-acknowledge"],
        ["Parameter yang dipantau", "Vibration saja", "Vibration (asli) + Temperature & Current (sintetis, disclosed)"],
        ["Filter parameter", "Tidak ada", "Ada, di Fleet Overview dan tab Sensors"],
        ["Fungsi temperature & current", "Cuma dipajang (tidak ada status/kesimpulan)", "Ada label Normal/Elevated per parameter + kalimat ringkasan otomatis"],
        ["Lihat data tanpa pindah halaman", "Tidak bisa", "Bisa, lewat tombol \"Quick view\""],
        ["Logo", "Tidak ada / tulisan biasa", "Logo Astra Otoparts asli"],
        ["Tampilan angka", "Font teks biasa", "Font monospace (JetBrains Mono), kesan presisi"],
        ["Background halaman", "Polos rata warna abu-abu", "Pattern heksagonal bergerak pelan + cahaya lembut"],
        ["Bug Acknowledge saat Play", "Ada (harus refresh halaman)", "Sudah diperbaiki dan dites ulang"],
        ["Perhitungan Health Score/Alarm/RUL", "Dari data vibrasi", "SAMA - tetap dari data vibrasi, tidak berubah"],
    ],
)

doc.add_page_break()

# ============================================================ 11. ASTRA
h1("12. Apakah Sudah Sesuai dengan Permintaan Astra?")
body(
    "Brief Case 2 dari Astra minta 6 hal utama (\"Key Features to Build\"). "
    "Berikut status masing-masing setelah update ini:"
)
status_table(
    ["Permintaan Astra", "Status", "Keterangan"],
    [
        ["Pengumpulan & monitoring data multi-parameter", "[OK] Sudah", "Vibration (asli) + Temperature/Current (sintetis) tampil bersamaan dan masing-masing punya status Normal/Elevated sendiri; Fleet Overview memantau 8 mesin sekaligus"],
        ["Model AI/ML deteksi pola degradasi", "[OK] Sudah", "Isolation Forest - persis contoh algoritma yang disebut di brief Astra"],
        ["Pengurangan false alarm (kombinasi parameter, bukan threshold tunggal)", "[OK] Sudah", "Titik terkuat proyek ini - aturan persistence + voting, sudah terbukti mengurangi alarm palsu drastis"],
        ["Estimasi Remaining Useful Life (RUL)", "[OK] Sudah", "Regresi eksponensial, sudah divalidasi dengan angka RMSE"],
        ["Dashboard real-time dengan Health Score", "[SEBAGIAN] Sebagian", "Dashboard-nya ada dan lengkap, tapi ini REPLAY data historis, bukan streaming langsung dari sensor sungguhan (karena memang belum ada sensor Astra yang terhubung)"],
        ["Sistem alert otomatis + rekomendasi tindakan", "[OK] Sudah", "Toast, banner critical + suara + acknowledge, plus tab berisi root cause dan saran tindakan"],
    ],
)
note_box(
    "Satu-satunya yang berstatus \"sebagian\" adalah soal \"real-time\" - dan itu "
    "bukan kekurangan yang disembunyikan, tapi konsekuensi wajar dari memakai "
    "dataset publik (bukan sensor Astra yang sungguhan terhubung). Ini sudah "
    "dijelaskan sejak proposal awal dan sebaiknya tetap disampaikan dengan jujur "
    "kalau ditanya."
)

doc.add_page_break()

# ============================================================ 12. DOSEN
h1("13. Apakah Sudah Sesuai dengan Feedback Dosen?")
status_table(
    ["Feedback Dosen", "Status", "Keterangan"],
    [
        ["Tambah menu \"Dashboard\"/high-level view untuk banyak mesin", "[OK] Sudah", "Halaman Fleet Overview, format sesuai contoh (\"Mesin X - Bearing Y → critical dalam Xh\")"],
        ["Notifikasi/alarm lebih mencolok, muncul berulang sampai ditindaklanjuti", "[OK] Sudah", "Banner berkedip + suara alarm loop + wajib klik Acknowledge"],
        ["Boleh dipertimbangkan data sintetis untuk parameter lain", "[OK] Sudah", "Temperature & Current, dibuat dengan logika fisika yang masuk akal, bukan random"],
        ["Filter berdasarkan parameter tambahan", "[OK] Sudah", "Dropdown \"Parameter focus\" + tab Sensors dengan pilihan tampilan"],
        ["Tampilan multi-mesin sekaligus (command center)", "[OK] Sudah", "Grid 8 kartu di Fleet Overview, dikelompokkan per Test"],
    ],
)
body(
    "Kesimpulannya: SEMUA poin yang diminta dosen sudah dikerjakan."
)

doc.add_page_break()

# ============================================================ 13. JUJUR
h1("14. Hal-Hal yang Harus Jujur Disampaikan (Kalau Ditanya)")
body(
    "Supaya tidak salah jawab atau terkesan menutup-nutupi saat presentasi/Q&A, "
    "ini poin-poin yang sebaiknya disampaikan apa adanya kalau ditanya:"
)
numbered(
    "seluruh pipeline deteksi (health score, alarm, RUL) HANYA dari data getaran "
    "(vibration) dari NASA IMS Bearing Dataset. Temperature dan current itu data "
    "buatan/simulasi, dibuat berdasarkan hasil perhitungan vibrasi yang sudah ada "
    "(dijelaskan detail di Bagian 5), BUKAN dari sensor sungguhan.",
    bold_lead="Data asli vs data buatan - ",
)
numbered(
    "dashboard ini adalah REPLAY (memutar ulang) data historis yang sudah pernah "
    "direkam, bukan aliran data langsung (live streaming) dari sensor yang "
    "sungguhan terpasang di mesin. Ini karena Astra belum menyediakan sensor "
    "yang terhubung.",
    bold_lead="Replay vs real-time - ",
)
numbered(
    "logo Astra Otoparts yang dipakai adalah file gambar resmi yang diberikan "
    "user sendiri, bukan hasil generate AI atau comot sembarangan dari internet.",
    bold_lead="Soal logo - ",
)
numbered(
    "tampilan yang lebih \"mewah\" (warna, animasi, font, pattern) itu murni "
    "soal tampilan/estetika - tidak mengubah sedikit pun logika perhitungan "
    "atau hasil analisis yang sudah divalidasi sebelumnya.",
    bold_lead="Soal tampilan baru - ",
)

doc.add_page_break()

# ============================================================ 14. GLOSSARY
h1("15. Daftar Istilah (Glossary)")
body("Penjelasan istilah-istilah yang muncul di dokumen ini, ditulis sesederhana mungkin:")

glossary = [
    ("Fleet Overview", "Halaman baru yang menampilkan status SEMUA mesin/bearing sekaligus dalam bentuk kartu-kartu, seperti ruang kontrol yang memantau banyak mesin bersamaan."),
    ("Health Score", "Angka 0-100 yang menunjukkan seberapa sehat kondisi bearing saat itu. 100 = sangat sehat, semakin kecil semakin rusak. Dihitung dari data getaran."),
    ("Alarm Level (Normal/Warning/Critical)", "Status kondisi mesin. Normal = aman, Warning = mulai ada tanda-tanda tidak beres (perlu diawasi lebih ketat), Critical = sudah parah, perlu tindakan segera."),
    ("RUL (Remaining Useful Life)", "Perkiraan \"berapa lama lagi mesin ini bisa jalan sebelum benar-benar rusak\", dalam satuan jam. Dihitung pakai model matematika dari tren kerusakan yang terdeteksi."),
    ("Vibration (Getaran)", "Data asli dari sensor getaran mesin - satu-satunya data ASLI yang tersedia di dataset NASA IMS Bearing yang dipakai proyek ini."),
    ("Temperature & Current (Suhu & Arus Listrik)", "Dua data TAMBAHAN yang dibuat/disimulasikan (bukan data sensor asli), untuk menunjukkan kesiapan dashboard memantau banyak parameter sekaligus, sesuai permintaan Astra."),
    ("Data Sintetis", "Istilah untuk \"data buatan/simulasi\" - bukan hasil pengukuran sensor sungguhan, tapi dihitung/diperkirakan dari data lain yang sudah ada, dengan logika yang masuk akal."),
    ("hi_norm / Health Index", "Angka internal (bukan yang ditampilkan langsung ke user) yang menunjukkan \"seberapa jauh kondisi bearing dari kondisi sehatnya\" - dipakai sebagai dasar untuk menghitung Health Score DAN sebagai dasar untuk membuat data suhu/arus listrik sintetis."),
    ("Fleet Clock (Fleet Time)", "Slider/penggeser yang menggerakkan SEMUA mesin ke titik waktu yang sebanding (dalam persen), meskipun panjang data tiap mesin berbeda-beda."),
    ("Modal / Dialog (Quick View)", "Jendela kecil yang muncul melayang DI ATAS halaman yang sedang dibuka, tanpa harus pindah halaman - seperti pop-up konfirmasi di aplikasi HP."),
    ("Acknowledge", "Tindakan menekan tombol untuk mengonfirmasi \"saya sudah lihat/sadar ada alarm ini\", supaya notifikasinya berhenti mengganggu."),
    ("Autoplay / Mode Play", "Fitur yang menggerakkan slider posisi waktu secara otomatis terus-menerus (seperti tombol Play di video), tanpa perlu digeser manual satu-satu."),
    ("Race Condition (istilah teknis pemrograman)", "Kondisi ketika dua proses yang berjalan hampir bersamaan \"berebut\" dan salah satunya jadi tidak sempat selesai dengan benar - ini yang menyebabkan bug tombol Acknowledge tidak bisa diklik saat mode Play jalan (dijelaskan detail di Bagian 9.1)."),
    ("CSS", "Bahasa pemrograman khusus untuk mengatur TAMPILAN (warna, ukuran, jarak, animasi) suatu halaman web, terpisah dari logika/perhitungan di baliknya."),
    ("Gradient (Gradasi)", "Perpaduan warna yang berubah secara halus dari satu warna ke warna lain (misalnya dari biru tua ke biru muda), dipakai di bar judul dashboard."),
    ("Shadow (Bayangan)", "Efek bayangan halus di sekitar kotak/kartu supaya terlihat \"terangkat\"/tiga dimensi, bukan menempel rata di layar."),
    ("Font Monospace", "Jenis huruf di mana semua karakter (termasuk angka) punya lebar yang sama persis - biasa dipakai di aplikasi teknis/coding karena kesannya presisi dan mudah dibandingkan angka per angka."),
    ("Animasi/Keyframe", "Pengaturan supaya suatu elemen di layar berubah/bergerak secara bertahap dan berulang (misalnya background yang bergeser pelan, atau badge yang berkedip)."),
    ("Base64", "Cara mengubah file gambar/suara menjadi teks panjang supaya bisa \"ditempel langsung\" di dalam kode halaman web, tanpa perlu file terpisah."),
    ("3-sigma Threshold (Batas 3-sigma)", "Garis batas statistik yang menandai \"nilai ini sudah jauh di luar kebiasaan normal\" - dipakai sebagai salah satu acuan untuk mendeteksi keanehan pada data sensor."),
    ("Badge Normal/Elevated", "Label kecil berwarna hijau (Normal) atau oranye (Elevated) di tiap grafik parameter pada tab Sensors - pembacaan sederhana yang membandingkan nilai parameter itu dengan batas normalnya sendiri, terpisah dari status alarm resmi (yang tetap dari data getaran)."),
]
for term, definition in glossary:
    p = doc.add_paragraph()
    r = p.add_run(term + " - ")
    r.bold = True
    r.font.color.rgb = BLUE
    p.add_run(definition)
    p.paragraph_format.space_after = Pt(8)

doc.add_page_break()

# ============================================================ PENUTUP
h1("Penutup")
body(
    "Semua perubahan di dokumen ini sudah dicoba langsung di browser (bukan cuma "
    "ditulis kodenya lalu diasumsikan jalan) - termasuk kasus bug Acknowledge yang "
    "sempat gagal, sampai benar-benar dipastikan sudah beres lewat percobaan "
    "ulang. Kalau ada pertanyaan lebih lanjut soal bagian mana pun di dokumen "
    "ini, atau butuh penjelasan yang lebih dalam untuk persiapan Q&A presentasi, "
    "tinggal ditanyakan kembali."
)

doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}")
