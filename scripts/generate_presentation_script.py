"""Generate the spoken narration script for the Week 2 progress presentation.

Pairs with docs/Week2_Progress_Presentation_Simple.pptx (9 slides) and the
live dashboard demo. Written to be read/rehearsed from, English narration,
short Indonesian prep notes in the Q&A backup section only.

Usage (from project root):
    python -m scripts.generate_presentation_script

Produces docs/Week2_Presentation_Script.docx
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "docs" / "Week2_Presentation_Script.docx"

NAVY = RGBColor(0x0B, 0x2A, 0x4A)
ACCENT = RGBColor(0x1F, 0x6F, 0xC4)
GRAY = RGBColor(0x55, 0x5B, 0x66)
GOOD = RGBColor(0x1E, 0x7B, 0x34)

doc = Document()

# base style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

for section in doc.sections:
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def h1(text, color=NAVY):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(20); r.font.bold = True; r.font.color.rgb = color
    p.space_after = Pt(6)
    return p


def h2(text, color=ACCENT):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(15); r.font.bold = True; r.font.color.rgb = color
    p.space_before = Pt(14); p.space_after = Pt(4)
    return p


def say(text, size=12):
    """Narration text - what the presenter actually says, quoted."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.italic = False
    p.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.5)
    return p


def note(text, color=GRAY):
    """Stage direction / presenter note, not spoken aloud."""
    p = doc.add_paragraph()
    r = p.add_run("» " + text)
    r.font.size = Pt(10.5); r.font.italic = True; r.font.color.rgb = color
    p.space_after = Pt(10)
    p.paragraph_format.left_indent = Cm(0.5)
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(11)
    return p


def hr():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ======================================================================
# Cover
# ======================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Predictive Maintenance Dashboard")
r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Week 2 Progress — Presentation Script")
r.font.size = Pt(16); r.font.color.rgb = ACCENT

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("Case 2 · AOP WINTEQ Bootcamp · Track 5 (Dashboard) · ~30 minutes")
r.font.size = Pt(11); r.font.color.rgb = GRAY

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run(
    "How to use this document: read it once fully before presenting. The italic lines "
    "marked » are stage directions (what to click/do), not things to say out loud. "
    "Everything else is narration — don't read it word for word, use it as your safety net."
)
r.font.size = Pt(10.5); r.font.color.rgb = GRAY; r.font.italic = True

doc.add_page_break()

# ======================================================================
# Timing overview
# ======================================================================
h1("Timing Overview (30-minute slot)")
table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Time"; hdr[1].text = "Section"; hdr[2].text = "Slide(s)"
for c in hdr:
    for p in c.paragraphs:
        for r in p.runs:
            r.font.bold = True

rows = [
    ("0:00 – 1:30", "Title + Problem & Goal", "1–2"),
    ("1:30 – 3:00", "Our Data (why NASA IMS)", "3"),
    ("3:00 – 5:30", "How It Works (plain language)", "4"),
    ("5:30 – 6:30", "Set up the live demo", "5"),
    ("6:30 – 18:00", "LIVE DEMO — one machine, healthy to failure", "(dashboard)"),
    ("18:00 – 21:00", "This Week: found & fixed a bug", "6"),
    ("21:00 – 23:00", "Honest limitations", "7"),
    ("23:00 – 25:00", "Next steps", "8"),
    ("25:00 – 30:00", "Thank you + Q&A buffer", "9"),
]
for t, sec, sl in rows:
    row = table.add_row().cells
    row[0].text = t; row[1].text = sec; row[2].text = sl

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("The live demo is the main event — if you're running short on time, cut from the closing sections, never from the demo.")
r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = NAVY

doc.add_page_break()

# ======================================================================
# Slide-by-slide narration
# ======================================================================
h1("Section 1 — Slide 1: Title")
say("“Good morning/afternoon. Today I'll show you the predictive maintenance "
    "dashboard we've built — and instead of walking you through slides of theory, "
    "I'm going to show you it working live, on one real machine, from healthy all "
    "the way to failure.”")

h1("Section 2 — Slide 2: The Problem")
say("“Right now, maintenance teams usually find out something is wrong only after "
    "it's already visible — vibration, noise, heat. By then the action window is short.”")
say("“The other issue is false alarms. If you set one threshold per sensor, you get "
    "so many noisy alerts that people start ignoring them — including the real ones.”")
say("“What we want: catch problems earlier, cut down false alarms, estimate how much "
    "time is left, and show all of it on one simple screen.”")

h1("Section 3 — Slide 3: Our Data")
say("“Astra's own sensor data isn't ready yet, so to prove the concept works, we used "
    "a public dataset that's the standard benchmark for exactly this problem: a real "
    "motor bearing, monitored with vibration sensors every 10 minutes, all the way "
    "from healthy to failure.”")
say("“One honest limitation up front: today's prototype only uses vibration. Real "
    "deployment will add temperature, current and pressure sensors — the dashboard "
    "is built to take those in later without a redesign, we just don't have that data "
    "yet.”")

h1("Section 4 — Slide 4: How It Works")
say("“Very briefly, in plain terms: every new sensor reading gets compared against what "
    "'healthy' looked like for this exact machine. Once several readings look abnormal "
    "together — not just one noisy blip — the system raises an alarm.”")
say("“Only after that alarm fires does it start estimating a countdown to failure. It "
    "never guesses before there's real evidence — that's a deliberate choice, not a "
    "limitation.”")
note("Do not read out the formulas even if you know them — keep this slide to the "
     "spoken lines above. If asked for details, that's what the Q&A backup notes at "
     "the end of this document are for.")

doc.add_page_break()

# ======================================================================
# LIVE DEMO — the main event
# ======================================================================
h1("Section 5 — LIVE DEMO (the main event, ~12 minutes)")
say("“Let's stop talking about it and just look at it. I'm switching to the live "
    "dashboard now.”")
note("Switch to the browser tab with the dashboard open (streamlit run app/dashboard.py). "
     "In the sidebar: Production run = 'Test 2 (dev run)', Machine/bearing = "
     "'Motor Conveyor - Bearing 1'. Click Reset so you start from position 0.")

h2("Checkpoint 1 — Healthy (position 0, t = 0.0h)")
say("“This is the machine at the very start of monitoring. Health score: 100. "
    "Status: Normal. No alarm, nothing to report — this is what 'boring and healthy' "
    "looks like, and it's what we want most of the time.”")
note("Point at the KPI row: Health Score gauge (100), Alarm Status (Normal).")

h2("Checkpoint 2 — First Warning (drag slider to ≈ position 537, t = 89.5h)")
say("“Now I'll jump forward. At around hour 89 into this machine's life, the status "
    "changes to Warning. Health score has only dropped slightly, to about 99 — "
    "barely visible by eye — but the system has already caught something: one "
    "vibration feature has been abnormal for several readings in a row, not just "
    "once.”")
say("“This is the false-alarm protection at work: it waited for a consistent pattern, "
    "not a single spike, before saying anything.”")
note("Point at the Alarm Status badge (Warning, orange) and read the 'Reason' text "
     "next to it out loud: 'Persistence: rms abnormal for 5 consecutive windows'.")

h2("Checkpoint 3 — Critical + prediction (drag slider to ≈ position 653, t = 108.8h)")
say("“Moving further ahead, about 19 hours later, the status escalates to Critical — "
    "three different vibration measurements are abnormal at the same time now, which "
    "is a much stronger signal than before.”")
say("“And this is the moment the system starts predicting: on the right, Remaining "
    "Useful Life now shows an estimate — about 4 days. That number only appears once "
    "there's real evidence of a problem, never before.”")
note("Point at the RUL KPI card (should read something like '4d 0h'). Optionally "
     "switch to the 'Alerts & Recommendation' tab and read the root cause + "
     "recommendation text out loud — it's written in plain maintenance language.")
say("“Notice the dashboard doesn't just say 'something's wrong' — it tells the "
    "maintenance team what's likely wrong (a bearing defect) and what to do about it "
    "(schedule replacement within the estimated window).”")

h2("Optional — Signal Detail tab (if time allows)")
say("“Just to show this isn't a black box: this chart is the actual vibration signal, "
    "and this second chart shows where in the frequency spectrum the damage signature "
    "shows up — that peak lines up with the known fault frequency for this exact "
    "bearing's geometry. It's physics-based, not guesswork.”")
note("Click 'Signal Detail' tab, point at the sharp peak in the envelope spectrum "
     "chart aligned with the 'BPFO' dashed line.")

h2("Wrap-up")
say("“So to summarize what you just watched: health score trending down, an alarm "
    "that only fires on a real, sustained pattern, and a time-remaining estimate that "
    "only appears once there's evidence — all on one screen, updating automatically.”")
note("Switch back to the slides.")

doc.add_page_break()

# ======================================================================
# Remaining slides
# ======================================================================
h1("Section 6 — Slide 6: This Week's Bug Fix")
say("“While preparing this demo, we noticed our own numbers on our longest test run "
    "looked wrong compared to an earlier draft. Instead of assuming either version "
    "was right, we investigated.”")
say("“We found that a single noisy sensor reading — one moment of coincidence — could "
    "trigger a false alarm, even while the machine was still perfectly healthy. On "
    "our longest test, that happened very early, and it threw off every prediction "
    "that came after it.”")
say("“The fix: an alarm now needs two consecutive abnormal readings before it counts "
    "as real, not just one. This chart shows the difference — the red line is where "
    "the old bug used to trigger, while the machine was still 100% healthy; the green "
    "line is where it triggers now, right as the health score actually starts to drop.”")
note("Point at the chart on this slide while saying the last line.")

h1("Section 7 — Slide 7: Honest Limitations")
say("“We want to be upfront about what this prototype does not do yet.”")
bullet("Vibration only — temperature, current and pressure come next, once real Astra data is available.")
bullet("The time-remaining number can still swing in the first hours right after an alarm fires, before enough data builds up.")
bullet("Tested on one operating condition. Real factory motors run under varying speed and load — that's future work.")

h1("Section 8 — Slide 8: Next Steps")
say("“For Week 3, our plan is to: stabilise the early-stage time-remaining prediction, "
    "prepare the dashboard to accept new sensor types once Astra's data is ready, "
    "test the same approach on other machines in the dataset, and start designing "
    "for live streaming data instead of replaying historical files.”")

h1("Section 9 — Slide 9: Thank You")
say("“Thank you — happy to take any questions.”")

doc.add_page_break()

# ======================================================================
# Q&A backup notes
# ======================================================================
h1("Q&A Backup Notes (not on slides — only use if asked)")
p = doc.add_paragraph()
r = p.add_run(
    "Keep answers short and simple first; only go deeper if the follow-up question "
    "asks for it. Indonesian notes in brackets are for your own prep, not to say out loud."
)
r.font.italic = True; r.font.size = Pt(10.5); r.font.color.rgb = GRAY

qa = [
    ("How does the system decide what's 'abnormal'?",
     "“It learns from the machine's own healthy period first — the first half of the "
     "run, before any problem exists. Anything that statistically stands out from that "
     "healthy baseline counts as abnormal.” "
     "[Ini Isolation Forest — model unsupervised, dilatih cuma di data sehat, threshold "
     "3-sigma dari skor data sehat.]"),
    ("Why not use a labelled/supervised model?",
     "“Because we don't have labels — nobody tags every sensor reading as 'healthy' or "
     "'faulty' in real time. We only know how each run ended. That's actually realistic: "
     "most real factory data looks the same way.” "
     "[Data cuma tau hasil akhir run, bukan label per-titik-waktu.]"),
    ("How is the remaining-time number calculated?",
     "“Once the alarm fires, we fit a curve to the trend and project forward to a "
     "failure point. We don't show the formula because it's not the point — what "
     "matters is that it only starts once there's real evidence, and it gets refit "
     "as new data comes in.” "
     "[Kalau didesak: HI(t)=p1*e^(p2t)+p3, EOL dari titik kurva itu nyentuh ambang "
     "ET, ET diambil dari run lain biar gak leakage.]"),
    ("Why did the two runs (Test 2 and Test 3) give different results before the fix?",
     "“Test 3 is 6.5× longer, and the bug we found only showed up clearly on that "
     "longer run — a coincidence was more likely to happen somewhere across a much "
     "longer recording. That's exactly why we test on a long, untouched holdout run: "
     "it exposes problems a short run wouldn't.”"),
    ("What about temperature, current, and pressure sensors from the brief?",
     "“The brief asks for those, and our architecture supports adding them — we just "
     "don't have that data yet since Astra hasn't shared it. This prototype proves "
     "the approach on vibration first.”"),
    ("Is this dataset realistic / good enough?",
     "“It's the standard public benchmark researchers use for exactly this kind of "
     "run-to-failure prediction problem — it's not our own choice made lightly, it's "
     "what the field itself uses when real plant data isn't available yet.”"),
    ("What happens if the alarm is wrong (false alarm)?",
     "“That's exactly what the persistence and voting rules are for — we already cut "
     "false alarms by removing single-reading noise before it counts. Our before/after "
     "comparison this week is a concrete example of catching and fixing exactly that "
     "kind of false trigger.”"),
]

for q, a in qa:
    p = doc.add_paragraph()
    r = p.add_run("Q: " + q)
    r.font.bold = True; r.font.size = Pt(12); r.font.color.rgb = NAVY
    p.space_before = Pt(10)
    p2 = doc.add_paragraph()
    r2 = p2.add_run("A: " + a)
    r2.font.size = Pt(11)
    p2.paragraph_format.left_indent = Cm(0.5)

OUT_PATH.parent.mkdir(exist_ok=True)
doc.save(str(OUT_PATH))
print("Saved:", OUT_PATH)
