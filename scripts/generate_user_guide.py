"""Generates docs/Panduan_Lengkap_Dashboard_Case2.docx - a full, step-by-step,
non-technical user guide to the Predictive Maintenance dashboard prototype.

Run:  python -m scripts.generate_user_guide   (from src/)
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm

OUT_PATH = Path(__file__).resolve().parents[2] / "docs" / "Panduan_Lengkap_Dashboard_Case2.docx"

NAVY = RGBColor(0x0F, 0x17, 0x2A)
BLUE = RGBColor(0x03, 0x69, 0xA1)
GREY = RGBColor(0x47, 0x55, 0x69)
GREEN = RGBColor(0x16, 0xA3, 0x4A)
RED = RGBColor(0xDC, 0x26, 0x26)

doc = Document()

# ---------------------------------------------------------------- base style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def set_margins():
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)


set_margins()


def h1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = NAVY
    return p


def h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = BLUE
    return p


def h3(text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = GREY
    return p


def body(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(8)
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def note_box(text, color=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("CATATAN: " + text)
    run.italic = True
    run.font.color.rgb = color
    return p


def simple_table(headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


# ============================================================ COVER
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("PANDUAN LENGKAP")
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Dashboard Predictive Maintenance  -  Induction Motor")
run.font.size = Pt(16)
run.font.color.rgb = BLUE

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub2.add_run("Case 2  -  Bootcamp AOP Winteq · PT Astra Otoparts Tbk")
run.font.size = Pt(12)
run.font.color.rgb = GREY

doc.add_paragraph()
note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = note.add_run(
    "Dokumen ini menjelaskan dashboard prototype dari awal sampai akhir: apa "
    "fungsinya, cara menjalankannya, cara membacanya, dan bagaimana \"otak\" "
    "di baliknya bekerja  -  ditulis dengan bahasa sesederhana mungkin."
)
run.italic = True
run.font.color.rgb = GREY

doc.add_page_break()

# ============================================================ DAFTAR ISI (manual)
h1("Daftar Isi")
toc_items = [
    "1. Apa Itu Dashboard Ini?",
    "2. Apakah Sudah Sesuai dengan yang Diminta Astra?",
    "3. Cara Menjalankan Dashboard (Langkah demi Langkah)",
    "4. Mengenal Data: Test 2, Test 3, dan Bearing 1-4",
    "5. Bagaimana \"Otak\" di Balik Dashboard Ini Bekerja",
    "6. Tur Lengkap Tampilan Dashboard",
    "   6.1 Bar Judul Paling Atas",
    "   6.2 Sidebar  -  Panel Kontrol di Kiri",
    "   6.3 Tiga Kartu Ringkasan (KPI)",
    "   6.4 Tab \"Trend\"",
    "   6.5 Tab \"Signal Detail\"",
    "   6.6 Tab \"Alerts & Recommendation\"",
    "   6.7 Catatan Metodologi di Paling Bawah",
    "7. Fungsi Play, Pause, Reset, dan Speed  -  Detail Lengkap",
    "8. Yang Sudah Selesai vs Yang Belum",
    "9. Pertanyaan yang Sering Muncul (FAQ)",
    "10. Daftar Istilah (Glossary)",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================ 1
h1("1. Apa Itu Dashboard Ini?")
body(
    "Dashboard ini adalah PROTOTYPE (contoh kerja) dari sistem yang diusulkan di "
    "Proposal Case 2: sebuah aplikasi yang memantau kondisi motor induksi (lewat "
    "bearing/laher-nya) dan mencoba MENEBAK lebih awal kapan motor itu akan rusak, "
    "sebelum benar-benar berhenti mendadak."
)
body(
    "Analoginya seperti alat pemantau detak jantung di rumah sakit: bukan cuma "
    "kasih tahu \"pasien sehat\" atau \"pasien sakit\", tapi juga kasih angka "
    "(health score), kasih peringatan dini kalau ada yang mencurigakan, dan "
    "coba memperkirakan \"berapa lama lagi sebelum kondisinya kritis\"."
)
bullet("Dijalankan secara lokal di komputer (belum terhubung ke internet/cloud).", bold_lead="Sifatnya: ")
bullet(
    "NASA IMS Bearing Dataset  -  data getaran motor asli dari Amerika yang sudah "
    "diuji sampai benar-benar rusak (run-to-failure). Dipakai karena Astra belum "
    "mengirim data motor asli, dan Miss sudah menyetujui pemakaian dataset ini "
    "untuk tahap prototype.",
    bold_lead="Data yang dipakai: ",
)
bullet(
    "Streamlit  -  dibuat dengan bahasa Python, jadi tampil sebagai halaman web "
    "yang dibuka lewat browser di alamat localhost:8501.",
    bold_lead="Teknologi: ",
)

# ============================================================ 2
h1("2. Apakah Sudah Sesuai dengan yang Diminta Astra?")
body(
    "Ya. Berikut perbandingan poin demi poin antara permintaan resmi Astra "
    "(dari dokumen Bootcamp Briefing, Use Case 2) dan apa yang sudah dibangun:"
)
simple_table(
    ["Yang Diminta Astra", "Status", "Keterangan"],
    [
        ["Monitoring multi-parameter", "Selesai", "RMS, kurtosis, crest factor, energi getaran BPFO/BPFI/BSF"],
        ["Deteksi pola degradasi berbasis AI", "Selesai", "Isolation Forest, dilatih hanya dari data motor sehat"],
        ["Pengurangan false alarm", "Selesai", "Aturan persistence + voting, terbukti mengurangi alarm palsu"],
        ["Estimasi Remaining Useful Life (RUL)", "Selesai", "Model eksponensial, hanya aktif setelah alarm pertama terkonfirmasi"],
        ["Health scoring real-time + dashboard", "Selesai", "Skor 0-100 + gauge visual di dashboard"],
        ["Automated alerts + rekomendasi", "Selesai", "Tab \"Alerts & Recommendation\" otomatis kasih saran tindakan"],
    ],
)
note_box(
    "Satu catatan jujur: \"multi-parameter\" di sini adalah multi-FITUR dari "
    "getaran (bukan gabungan sensor arus + suhu + tekanan beneran), karena "
    "dataset NASA IMS memang cuma merekam getaran. Sensor arus dan suhu masuk "
    "rencana tahap pilot lapangan (bukan bagian prototype minggu ini)  -  ini "
    "sudah dijelaskan secara terbuka di Proposal juga, bukan sesuatu yang "
    "disembunyikan."
)
body("Dan dibandingkan dengan feedback Miss saat progress presentation:")
simple_table(
    ["Feedback Miss", "Status"],
    [
        ["Presentasi jangan terlalu banyak teori/rumus", "Dashboard tidak menampilkan rumus  -  hanya gauge, badge, dan grafik"],
        ["Dataset NASA IMS boleh dipakai", "Dipakai sesuai arahan"],
        ["Pendekatan unsupervised (Isolation Forest, 3 standar deviasi) benar", "Dipakai persis seperti itu"],
        ["Ingin lihat 1 mesin: alarm & RUL saat kondisi turun di bawah ambang", "Ada di tab Trend  -  grafik anomaly score vs garis ambang 3-sigma"],
    ],
)

# ============================================================ 3
h1("3. Cara Menjalankan Dashboard (Langkah demi Langkah)")
body("Ikuti urutan ini dari atas ke bawah, jangan ada yang dilewat:")
numbered("Buka Terminal / PowerShell / Command Prompt di komputer kamu.")
numbered("Pindah ke folder project dengan mengetik (lalu tekan Enter):")
p = doc.add_paragraph("cd C:\\Documents\\COMPETITION\\Astra_Case2\\src")
p.style = doc.styles["Intense Quote"]
numbered("Jalankan perintah berikut (lalu tekan Enter):")
p = doc.add_paragraph("python -m streamlit run app/dashboard.py")
p.style = doc.styles["Intense Quote"]
numbered(
    "Tunggu beberapa detik sampai muncul tulisan \"Local URL: "
    "http://localhost:8501\" di terminal."
)
numbered("Buka browser (Chrome/Edge), lalu ketik di address bar: localhost:8501")
numbered(
    "Dashboard akan langsung terbuka. JANGAN tutup jendela terminal tadi  -  itu "
    "\"mesin\" yang membuat dashboard tetap hidup. Kalau terminal ditutup, "
    "dashboard ikut mati."
)
numbered("Kalau mau berhenti: klik di jendela terminal, lalu tekan Ctrl + C.")
note_box(
    "Kalau muncul pesan \"port sudah dipakai\" atau semacamnya, kemungkinan "
    "dashboard sudah berjalan dari sesi lain  -  coba langsung buka "
    "localhost:8501 di browser dulu sebelum menjalankan ulang perintahnya."
)

# ============================================================ 4
h1("4. Mengenal Data: Test 2, Test 3, dan Bearing 1-4")
body(
    "Ini bagian yang paling sering bikin bingung, jadi dijelaskan pelan-pelan "
    "dari dasar."
)
h2("4.1 Apa itu \"Bearing\"?")
body(
    "Bearing adalah laher/bantalan  -  komponen di dalam motor yang membuat poros "
    "bisa berputar dengan mulus. Kalau bearing rusak, biasanya itu adalah "
    "penyebab kegagalan motor yang paling sering terjadi di industri (menurut "
    "riset yang dipakai tim ini, sekitar 41-42% dari semua kegagalan motor "
    "induksi berasal dari bearing yang rusak)."
)
h2("4.2 Kenapa ada 4 Bearing?")
body(
    "Di alat uji (test rig) yang menghasilkan dataset ini, ada SATU poros yang "
    "berputar, dan di sepanjang poros itu dipasang EMPAT bearing sekaligus, "
    "masing-masing direkam getarannya lewat sensor terpisah:"
)
body("[Bearing 1]  -  [Bearing 2]  -  [Bearing 3]  -  [Bearing 4]  -  semuanya di 1 poros yang sama, digerakkan 1 motor.", italic=True)
h2("4.3 Kenapa ada Test 2 dan Test 3?")
body(
    "\"Test 2\" dan \"Test 3\" adalah DUA REKAMAN PERCOBAAN YANG BERBEDA (dari "
    "3 yang tersedia dari NASA, tapi cuma 2 yang dipakai di prototype ini). "
    "Di kedua rekaman itu, ke-4 bearing dijalankan terus sampai salah satu "
    "bearingnya benar-benar rusak."
)
simple_table(
    ["", "Test 2 (dijadikan run pengembangan)", "Test 3 (dijadikan run holdout/ujian)"],
    [
        ["Lama rekaman", "163,8 jam (~6,8 hari)", "1073,3 jam (~44,7 hari)"],
        ["Jumlah data getaran", "984 potongan rekaman", "6.324 potongan rekaman"],
        ["Bearing mana yang rusak", "Bearing 1", "Bearing 3"],
        ["Jenis kerusakan", "Outer race (cincin luar bearing)", "Outer race (cincin luar bearing)"],
        ["Bearing 2, 3, 4 lainnya (Test 2) / 1, 2, 4 (Test 3)", "Tetap sehat sepanjang rekaman", "Tetap sehat sepanjang rekaman"],
    ],
)
note_box(
    "\"Holdout\" artinya: Test 3 sengaja TIDAK dilihat/diintip sama sekali "
    "sampai sistemnya benar-benar selesai dibangun dari Test 2  -  supaya hasil "
    "ujiannya jujur, bukan sistem yang \"udah lihat bocorannya duluan\". Ini "
    "prinsip penting yang juga dipakai di Proposal."
)
h2("4.4 Kesimpulan Test x Bearing")
body(
    "Karena ada 2 rekaman x 4 bearing, totalnya ada 8 kombinasi \"bearing-run\" "
    "yang masing-masing dianalisis terpisah dan disimpan sebagai 1 file data "
    "sendiri-sendiri. Di dashboard, kamu bisa pilih salah satu dari 8 "
    "kombinasi ini lewat sidebar."
)

# ============================================================ 5
h1("5. Bagaimana \"Otak\" di Balik Dashboard Ini Bekerja")
body(
    "Supaya kebayang, ini analoginya seperti proses pemeriksaan kesehatan "
    "berlapis. Data getaran mentah diproses lewat beberapa \"tahap pemeriksaan\" "
    "berurutan sebelum akhirnya muncul jadi angka-angka yang gampang dibaca di "
    "dashboard:"
)
h2("Tahap 1  -  Ambil Ciri-Ciri dari Getaran (Feature Extraction)")
body(
    "Setiap potongan rekaman getaran (1 detik, direkam 20.000 kali per detik) "
    "diringkas jadi beberapa angka ciri khas, seperti dokter yang mengukur "
    "detak jantung, tekanan darah, dan suhu tubuh  -  bukan menyimpan mentah "
    "rekaman EKG-nya:"
)
bullet("Seberapa besar/kuat getarannya secara keseluruhan.", bold_lead="RMS  -  ")
bullet(
    "Seberapa \"tajam/mengejutkan\" pola getarannya. Motor sehat nilainya "
    "sekitar 3; kalau ada benturan/cacat di bearing, angka ini naik drastis.",
    bold_lead="Kurtosis  -  ",
)
bullet(
    "Energi getaran di frekuensi tertentu yang secara fisika PASTI muncul "
    "kalau ada cacat di bagian tertentu bearing (dihitung dari ukuran fisik "
    "bearing-nya, bukan tebakan).",
    bold_lead="Energi BPFO / BPFI / BSF  -  ",
)
h2("Tahap 2  -  Deteksi Anomali (Isolation Forest)")
body(
    "Sistem \"belajar\" seperti apa pola normal HANYA dari data di masa motor "
    "masih sehat (separuh awal rekaman). Setelah itu, tiap titik waktu baru "
    "dibandingkan: seberapa \"aneh\" dia dibanding pola sehat yang sudah "
    "dipelajari? Kalau keanehannya melewati batas statistik (3 standar "
    "deviasi dari kondisi sehat), itu dianggap sinyal awal masalah."
)
h2("Tahap 3  -  Health Index / Skor Kesehatan")
body(
    "Beberapa ciri getaran tadi digabung jadi SATU angka tunggal (0-100) "
    "yang gampang dibaca siapa saja  -  100 berarti sehat sempurna, mendekati "
    "0 berarti mendekati gagal. Teknik penggabungannya (PCA) otomatis "
    "menentukan ciri mana yang paling penting dari datanya sendiri, bukan "
    "ditentukan asal tebak oleh manusia."
)
h2("Tahap 4  -  Logika Alarm (Persistence + Voting)")
body(
    "Supaya tidak False alarm terus-menerus, alarm baru benar-benar "
    "\"berbunyi\" kalau salah satu dari dua syarat ini terpenuhi:"
)
bullet(
    "Kalau 2 atau lebih ciri getaran SAMA-SAMA abnormal di waktu yang sama, "
    "alarm langsung berbunyi (dianggap bukti kuat).",
    bold_lead="Voting  -  ",
)
bullet(
    "Kalau cuma 1 ciri yang abnormal, alarm baru berbunyi setelah kondisi "
    "itu bertahan beberapa kali berturut-turut (supaya lonjakan sesaat/noise "
    "tidak langsung dianggap masalah serius).",
    bold_lead="Persistence  -  ",
)
h2("Tahap 5  -  Estimasi RUL (Remaining Useful Life)")
body(
    "Begitu alarm pertama kali \"resmi\" berbunyi (lolos persistence+voting), "
    "barulah sistem mulai menghitung perkiraan \"berapa jam lagi sampai "
    "gagal\", dengan mencocokkan tren penurunan Health Index ke sebuah kurva "
    "matematis (kurva eksponensial), dan terus dihitung ulang setiap ada data "
    "baru. SEBELUM alarm pertama berbunyi, sistem sengaja TIDAK berani "
    "menebak angka RUL  -  statusnya cuma \"monitoring, belum stabil\", karena "
    "menebak durasi kegagalan sebelum ada tanda-tanda nyata itu tidak jujur "
    "secara statistik."
)
note_box(
    "Batas kegagalan (\"seberapa rendah skor sampai dianggap benar-benar "
    "gagal\") diambil dari REKAMAN LAIN yang bearingnya sudah benar-benar "
    "rusak  -  bukan dari rekaman yang sedang dianalisis itu sendiri. Ini "
    "supaya sistemnya tidak \"nyontek\" masa depannya sendiri."
)

# ============================================================ 6
h1("6. Tur Lengkap Tampilan Dashboard")
body("Sekarang kita jalan-jalan keliling tampilan dashboard-nya, dari atas ke bawah.")

h2("6.1 Bar Judul Paling Atas")
body(
    "Bagian biru gelap paling atas menampilkan judul dashboard, nama "
    "perusahaan, dan di kanan ada info bearing/run yang sedang dipilih plus "
    "posisi waktu saat ini (dalam jam sejak rekaman dimulai)."
)

h2("6.2 Sidebar  -  Panel Kontrol di Kiri")
bullet(
    "Pilih mau lihat Test 2 (dev run) atau Test 3 (holdout run).",
    bold_lead="Production run  -  ",
)
bullet(
    "Pilih mau lihat bearing yang mana (1-4). Bearing yang memang diketahui "
    "rusak ditandai \"(monitored - known failure)\".",
    bold_lead="Machine / bearing  -  ",
)
bullet(
    "Slider yang menunjukkan \"kita sedang melihat titik waktu ke berapa\" "
    "dari seluruh rekaman. Bisa digeser manual, atau berjalan otomatis kalau "
    "Play ditekan (dijelaskan detail di Bab 7).",
    bold_lead="Snapshot position  -  ",
)
bullet("Tombol Play/Pause dan Reset  -  dijelaskan lengkap di Bab 7.", bold_lead="Play / Pause / Reset  -  ")
bullet(
    "Seberapa cepat lompatan waktunya saat mode Play aktif (1x paling "
    "pelan, 64x paling cepat).",
    bold_lead="Playback speed  -  ",
)
bullet(
    "Info kecil di bawah sidebar yang menyebutkan sumber datanya (NASA IMS "
    "Bearing Dataset)  -  supaya transparan datanya dari mana.",
    bold_lead="Keterangan sumber data  -  ",
)

h2("6.3 Tiga Kartu Ringkasan (KPI)")
body("Tiga kartu ini adalah \"ringkasan sekilas\" kondisi motor di titik waktu yang sedang dilihat:")
h3("Kartu 1  -  Health Score")
body(
    "Jarum/gauge melengkung dari 0 sampai 100. Semakin ke kanan (mendekati "
    "100) semakin sehat. Ada 3 warna latar: merah (0-40, kritis), kuning "
    "(40-70, waspada), hijau (70-100, sehat)."
)
h3("Kartu 2  -  Alarm Status")
body(
    "Menampilkan status saat ini: NORMAL (hijau), WARNING (kuning), atau "
    "CRITICAL (merah), lengkap dengan ikon dan alasan singkat kenapa status "
    "itu muncul (misalnya \"Voting: 2 parameter abnormal bersamaan\")."
)
h3("Kartu 3  -  Remaining Useful Life (RUL)")
body(
    "Kalau sistem sudah cukup yakin, di sini muncul perkiraan \"sisa berapa "
    "jam lagi sebelum gagal\" (contoh: \"3d 5h\" berarti 3 hari 5 jam). Kalau "
    "belum cukup data/belum ada alarm resmi, tulisannya \"Monitoring\"  -  "
    "artinya sistem sengaja belum berani menebak."
)

h2("6.4 Tab \"Trend\"")
body("Berisi 2 grafik berdampingan:")
bullet(
    "Garis biru menunjukkan Health Score dari waktu ke waktu (dari awal "
    "rekaman sampai posisi yang sedang dilihat). Titik-titik merah menandai "
    "kapan alarm sedang aktif. Ada garis putus-putus oranye menandai kapan "
    "\"trigger\" pertama terjadi (mulai dihitungnya RUL).",
    bold_lead="Grafik kiri (Health Score over time)  -  ",
)
bullet(
    "Ini yang PALING SESUAI dengan permintaan Miss: menunjukkan kapan "
    "kondisi motor \"turun\" dan melewati garis putus-putus merah (garis "
    "ambang 3 standar deviasi). Begitu garis abu-abu (skor anomali) melewati "
    "garis merah itu, itulah momen sistem mulai curiga ada masalah.",
    bold_lead="Grafik kanan (Anomaly score vs garis ambang 3-sigma)  -  ",
)

h2("6.5 Tab \"Signal Detail\"")
body("Berisi 2 grafik yang menunjukkan \"bukti mentah\" di balik keputusan sistem:")
bullet(
    "Potongan sinyal getaran asli (2000 titik pertama dari rekaman 1 detik) "
    "pada waktu yang sedang dilihat  -  ini data paling mentah, belum diolah "
    "sama sekali.",
    bold_lead="Grafik kiri (Raw vibration snapshot)  -  ",
)
bullet(
    "Menunjukkan di frekuensi berapa energi getaran itu paling besar, "
    "dengan garis putus-putus merah menandai frekuensi BPFO/BPFI/BSF (yang "
    "artinya cacat di bagian tertentu bearing). Kalau garis grafiknya "
    "\"memuncak tinggi\" tepat di salah satu garis merah itu, itu bukti "
    "fisik nyata bahwa memang ada kerusakan di bagian tersebut  -  bukan "
    "cuma tebakan statistik.",
    bold_lead="Grafik kanan (Envelope spectrum)  -  ",
)

h2("6.6 Tab \"Alerts & Recommendation\"")
body("Ringkasan dalam bentuk teks (bukan grafik), berisi:")
bullet("Status alarm saat ini beserta alasannya.", bold_lead="Status & Alasan  -  ")
bullet(
    "Tebakan sistem soal bagian mana yang bermasalah (untuk data ini, kalau "
    "memang bearing yang dipantau adalah bearing yang gagal, jawabannya "
    "\"Bearing outer race defect\").",
    bold_lead="Likely root cause  -  ",
)
bullet(
    "Saran tindakan otomatis, isinya menyesuaikan status: kalau CRITICAL, "
    "sarannya \"jadwalkan penggantian bearing\"; kalau WARNING, sarannya "
    "\"tingkatkan frekuensi pemantauan\"; kalau NORMAL, sarannya \"tidak "
    "perlu tindakan\".",
    bold_lead="Maintenance recommendation  -  ",
)

h2("6.7 Catatan Metodologi di Paling Bawah")
body(
    "Baris kecil di bagian paling bawah halaman, isinya ringkasan metode "
    "yang dipakai (Isolation Forest, PCA Health Index, RUL eksponensial, "
    "persistence+voting)  -  supaya siapa pun yang lihat dashboard-nya "
    "langsung tahu dasar ilmiahnya tanpa harus buka dokumen proposal."
)

# ============================================================ 7
h1("7. Fungsi Play, Pause, Reset, dan Speed  -  Detail Lengkap")
h2("Play")
body(
    "Kalau ditekan, dashboard akan \"memutar ulang\" data secara otomatis "
    "berjalan maju  -  persis seperti menonton rekaman CCTV time-lapse. Setiap "
    "beberapa saat, posisi waktu bergeser sendiri ke titik berikutnya, dan "
    "SEMUA yang ada di layar ikut update otomatis mengikuti posisi baru itu: "
    "Health Score naik-turun, Alarm Status bisa berubah dari NORMAL ke "
    "WARNING ke CRITICAL, angka RUL ikut dihitung ulang, grafik Trend "
    "bertambah panjang, dan sinyal di tab Signal Detail ikut berganti sesuai "
    "titik waktu yang baru. Tombolnya akan berubah tulisan jadi \"Pause\" "
    "selama mode ini aktif."
)
h2("Pause")
body(
    "Menghentikan proses \"Play\" tadi di posisi waktu yang sedang "
    "ditampilkan, supaya bisa diamati dengan tenang atau di-screenshot untuk "
    "keperluan presentasi."
)
h2("Reset")
body(
    "Mengembalikan posisi waktu ke titik \"default\" yang sudah dipilihkan "
    "otomatis oleh sistem (yaitu sesaat setelah alarm pertama kali resmi "
    "berbunyi)  -  supaya setiap ganti pilihan bearing/run, tidak perlu geser "
    "manual dari nol lagi untuk langsung melihat momen pentingnya."
)
h2("Playback speed")
body(
    "Mengatur seberapa jauh lompatan tiap langkah saat mode Play aktif: 1x "
    "paling pelan (lompat 1 titik data setiap kali), 64x paling cepat "
    "(lompat 64 titik data setiap kali)  -  cocok dipakai kalau ingin cepat "
    "melihat \"cerita penuh\" dari sehat sampai gagal tanpa menunggu lama."
)

# ============================================================ 8
h1("8. Yang Sudah Selesai vs Yang Belum")
h2("Sudah selesai dan berfungsi")
bullet("Ekstraksi ciri getaran dari data mentah (RMS, kurtosis, energi BPFO/BPFI/BSF).")
bullet("Deteksi anomali dengan Isolation Forest, ambang 3 standar deviasi.")
bullet("Health Index berbasis PCA, skala 0-100.")
bullet("Logika alarm persistence + voting, terbukti mengurangi jumlah alarm palsu.")
bullet("Estimasi RUL dengan model eksponensial, hanya aktif setelah alarm resmi.")
bullet("Dashboard interaktif lengkap dengan mode replay (Play/Pause/Reset).")
h2("Belum termasuk dalam prototype ini (rencana tahap berikutnya)")
bullet("Data getaran/arus/suhu asli dari motor Astra (masih memakai data NASA IMS).")
bullet("Koneksi langsung/live ke sistem PLC/SCADA pabrik (rencana fase pilot lapangan).")
bullet("Sensor arus 3 fasa dan suhu winding (dataset ini hanya punya sensor getaran).")
bullet("Deteksi jenis kegagalan lain selain bearing (rotor bar, stator winding, eccentricity).")

# ============================================================ 9
h1("9. Pertanyaan yang Sering Muncul (FAQ)")
h3("Kenapa health score kadang naik lagi padahal tadinya turun terus?")
body(
    "Ini bukan bug  -  data getaran asli memang tidak selalu turun mulus, "
    "kadang ada fluktuasi wajar (noise) sebelum akhirnya benar-benar jatuh "
    "menjelang kegagalan. Proposal Case 2 sendiri secara terbuka mengakui "
    "sifat data yang \"berosilasi\" ini sebagai salah satu keterbatasan yang "
    "diketahui."
)
h3("Kenapa Test 3 datanya jauh lebih banyak daripada Test 2?")
body(
    "Karena rekaman Test 3 memang berlangsung jauh lebih lama (1073 jam vs "
    "163,8 jam) sebelum bearingnya benar-benar rusak  -  bukan karena sensor "
    "yang berbeda, cuma durasi pengujiannya yang berbeda."
)
h3("Apakah angka RUL-nya bisa dipercaya 100%?")
body(
    "Tidak sepenuhnya, dan itu memang jujur diakui. Sebelum alarm resmi "
    "berbunyi, sistem sengaja TIDAK menampilkan angka RUL sama sekali "
    "(statusnya \"Monitoring\") justru supaya tidak menyesatkan. Setelah "
    "alarm berbunyi, angkanya adalah estimasi terbaik berdasarkan tren yang "
    "ada, tapi seperti prediksi cuaca  -  bisa berubah seiring data baru masuk."
)
h3("Kalau saya ganti pilihan bearing/run, apakah harus tunggu proses hitung ulang?")
body(
    "Tidak. Semua perhitungan berat (ekstraksi fitur, model, RUL) sudah "
    "dihitung DI AWAL lewat proses \"precompute\" dan disimpan ke file. "
    "Dashboard hanya menampilkan hasil yang sudah jadi, jadi berpindah "
    "bearing/run terasa instan."
)

# ============================================================ 10
h1("10. Daftar Istilah (Glossary)")
simple_table(
    ["Istilah", "Artinya (bahasa sederhana)"],
    [
        ["Bearing", "Laher/bantalan yang membuat poros motor berputar mulus"],
        ["Health Score", "Skor kesehatan 0-100, 100 = sehat sempurna, 0 = gagal"],
        ["Alarm / Alert", "Peringatan bahwa kondisi motor mencurigakan/bermasalah"],
        ["RUL (Remaining Useful Life)", "Perkiraan sisa waktu sebelum motor benar-benar gagal"],
        ["Isolation Forest", "Metode AI yang belajar pola \"sehat\" lalu mendeteksi apa pun yang menyimpang jauh dari pola itu"],
        ["3-sigma / 3 standar deviasi", "Batas statistik: kalau suatu nilai jauh berbeda (3x lipat dari sebaran normal) dari kondisi sehat, dianggap tidak wajar"],
        ["Persistence", "Aturan \"harus terjadi berulang kali dulu baru dianggap serius\", supaya lonjakan sesaat tidak langsung dianggap alarm"],
        ["Voting", "Aturan \"kalau 2 tanda mencurigakan muncul bersamaan, langsung dianggap serius\""],
        ["BPFO / BPFI / BSF", "Nama-nama frekuensi getaran tertentu yang secara fisika pasti muncul kalau ada cacat di bagian tertentu bearing"],
        ["Envelope spectrum", "Cara membaca sinyal getaran untuk menonjolkan pola benturan kecil akibat cacat bearing yang biasanya tersembunyi"],
        ["Run-to-failure", "Rekaman data yang berlangsung terus sampai komponennya benar-benar rusak total"],
        ["Holdout", "Data yang sengaja disimpan/tidak diintip sampai sistem selesai dibangun, dipakai khusus untuk menguji kejujuran hasil"],
    ],
)

doc.add_page_break()
end = doc.add_paragraph()
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = end.add_run(
    " -  Selesai  - \nDokumen ini dibuat untuk mendampingi dashboard prototype "
    "Case 2 Predictive Maintenance, PT Astra Otoparts Tbk / WINTEQ."
)
run.italic = True
run.font.color.rgb = GREY

doc.save(str(OUT_PATH))
print(f"Saved: {OUT_PATH}")
