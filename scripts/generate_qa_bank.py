"""Generate a comprehensive anticipated-questions Q&A bank for the Astra
Case 2 bootcamp project (not a script — a reference document to study).

Covers the whole project: business context, dataset, features, anomaly
detection, health score/RUL, false-alarm reduction, validation, honesty/
limitations, the dashboard itself, roadmap, and tricky conceptual questions.

Usage (from project root):
    python -m scripts.generate_qa_bank

Produces docs/Anticipated_Questions_QA_Bank.docx
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "docs" / "Anticipated_Questions_QA_Bank.docx"

NAVY = RGBColor(0x0B, 0x2A, 0x4A)
ACCENT = RGBColor(0x03, 0x69, 0xA1)
GRAY = RGBColor(0x55, 0x5B, 0x66)
SPOKEN = RGBColor(0x11, 0x14, 0x18)
ID_COLOR = RGBColor(0x0F, 0x76, 0x4A)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

for section in doc.sections:
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)


def h1(text, color=NAVY, size=24):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.bold = True; r.font.color.rgb = color
    p.space_after = Pt(8)
    return p


def h2(text, num, color=ACCENT):
    p = doc.add_paragraph()
    r = p.add_run(f"{num}. {text}")
    r.font.size = Pt(16); r.font.bold = True; r.font.color.rgb = color
    p.space_before = Pt(20); p.space_after = Pt(8)
    return p


def qa(question, answer_en, answer_id, qnum):
    p = doc.add_paragraph()
    r = p.add_run(f"Q{qnum}. {question}")
    r.font.bold = True; r.font.size = Pt(12); r.font.color.rgb = NAVY
    p.space_before = Pt(10); p.space_after = Pt(3)

    p2 = doc.add_paragraph()
    r2 = p2.add_run(answer_en)
    r2.font.size = Pt(11.5); r2.font.color.rgb = SPOKEN
    p2.paragraph_format.left_indent = Cm(0.4)
    p2.space_after = Pt(2)

    p3 = doc.add_paragraph()
    r3 = p3.add_run(answer_id)
    r3.font.size = Pt(10); r3.font.italic = True; r3.font.color.rgb = ID_COLOR
    p3.paragraph_format.left_indent = Cm(0.4)
    p3.space_after = Pt(6)


def hr():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ======================================================================
# Cover + table of contents
# ======================================================================
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Anticipated Questions — Q&A Bank")
r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = NAVY

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Case 2 · AOP WINTEQ Bootcamp · Predictive Maintenance Dashboard")
r.font.size = Pt(14); r.font.color.rgb = ACCENT

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("This is a study document, not a script. ")
r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = NAVY
r2 = p.add_run(
    "Don't memorize word for word — read each answer a few times, understand "
    "the idea, then answer in your own words if asked. Black text is the "
    "English answer. Green italic text is just the Indonesian meaning, for "
    "your own understanding."
)
r2.font.size = Pt(11); r2.font.color.rgb = GRAY

p = doc.add_paragraph()
r = p.add_run(
    "Dokumen ini buat belajar, bukan script. Gak usah dihafal kata per kata "
    "— baca beberapa kali, paham intinya, terus jawab pakai kata-kata "
    "sendiri kalau ditanya. Teks hitam itu jawaban Inggrisnya. Teks hijau "
    "miring itu artinya doang, buat paham sendiri."
)
r.font.size = Pt(10); r.font.italic = True; r.font.color.rgb = ID_COLOR
p.space_after = Pt(14)

h1("Contents", size=16)
sections = [
    "A. Business Context & WINTEQ",
    "B. The Dataset (NASA IMS Bearing)",
    "C. Features (RMS, kurtosis, envelope, BPFO/BPFI/BSF/FTF)",
    "D. Anomaly Detection (Isolation Forest)",
    "E. Health Score & Remaining Useful Life (RUL)",
    "F. False-Alarm Reduction",
    "G. Validation & This Week's Bug Fix",
    "H. Limitations & Honesty",
    "I. The Dashboard / Prototype Itself",
    "J. Roadmap & Future Work",
    "K. Team & Process",
    "L. Tricky / Conceptual Questions",
]
for s in sections:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(s)
    r.font.size = Pt(11.5); r.font.color.rgb = NAVY

doc.add_page_break()

# ======================================================================
# A. Business Context
# ======================================================================
h2("Business Context & WINTEQ", "A")

qa("What is WINTEQ and how does it relate to Astra Otoparts?",
   "WINTEQ is an engineering business unit under PT Astra Otoparts Tbk, started in 2006, "
   "focused on engineering, system integration, and factory automation.",
   "WINTEQ itu unit bisnis engineering di bawah PT Astra Otoparts Tbk, berdiri 2006, fokus "
   "ke engineering, integrasi sistem, dan otomasi pabrik.", 1)

qa("What business problem is this project solving?",
   "Right now, maintenance teams only find out something's wrong after it's already visible. "
   "And single-sensor alarms are so noisy that people start ignoring them, including the real ones.",
   "Sekarang, tim maintenance baru tau ada masalah setelah keliatan jelas. Dan alarm dari 1 "
   "sensor doang itu berisik banget, orang jadi cuek, termasuk pas alarm-nya beneran.", 2)

qa("What's the goal of this specific use case (Use Case 2)?",
   "Build an AI-based predictive maintenance system: monitor multiple parameters together, "
   "detect problems early, reduce false alarms, estimate remaining time before failure, and "
   "show all of it on one dashboard with automatic alerts and recommendations.",
   "Bikin sistem predictive maintenance berbasis AI: monitor banyak parameter sekaligus, "
   "deteksi masalah dari awal, kurangi false alarm, kasih estimasi sisa waktu sebelum gagal, "
   "dan tampilin semua di 1 dashboard dengan alert & rekomendasi otomatis.", 3)

qa("Why does this matter specifically for Astra?",
   "Induction motors run conveyors, pumps, compressors, and stamping or CNC machines across "
   "their plants — battery, piston, brake, rubber-plastic. Unplanned downtime breaks their "
   "just-in-time production chain with OEM customers.",
   "Motor induksi itu ngegerakin conveyor, pompa, kompresor, mesin stamping/CNC di "
   "pabrik-pabrik mereka. Downtime mendadak itu ngerusak rantai produksi just-in-time mereka "
   "ke customer OEM.", 4)

doc.add_page_break()

# ======================================================================
# B. Dataset
# ======================================================================
h2("The Dataset (NASA IMS Bearing)", "B")

qa("Why didn't you use real Astra sensor data?",
   "Because it's not available yet at this stage of the bootcamp.",
   "Karena datanya belum tersedia di tahap bootcamp sekarang.", 5)

qa("What dataset did you use instead, and why this one?",
   "The NASA IMS Bearing dataset — a public run-to-failure vibration dataset. We picked it "
   "because it's one of the only public datasets that's a real run-to-failure recording, not "
   "just fault classification, so it actually supports building a realistic RUL model.",
   "Dataset NASA IMS Bearing — dataset getaran run-to-failure yang publik. Kita pilih ini "
   "karena ini salah satu dataset publik langka yang beneran rekaman run-to-failure, bukan "
   "cuma klasifikasi fault, jadi beneran bisa dipakai bikin model RUL yang realistis.", 6)

qa("What's actually in this dataset?",
   "Four bearings on one shaft, spinning at a constant 2000 RPM. Vibration is recorded every "
   "10 minutes until failure, 20,000 samples per second, one-second snapshots each time.",
   "Empat bearing di satu poros, berputar konstan di 2000 RPM. Getaran direkam tiap 10 menit "
   "sampai gagal, 20.000 sampel per detik, tiap rekaman 1 detik.", 7)

qa("How many test runs did you use, and why two?",
   "Test 2 (163.8 hours, bearing 1 fails outer race) as the development run we used to build "
   "the system. Test 3 (1073.3 hours, 6.5x longer, bearing 3 fails outer race) as a completely "
   "untouched holdout, only checked at the very end, to prove the system isn't just tuned to "
   "fit one recording.",
   "Test 2 (163,8 jam, bearing 1 gagal outer race) sebagai run pengembangan yang kita pakai "
   "bikin sistemnya. Test 3 (1073,3 jam, 6,5x lebih panjang, bearing 3 gagal outer race) "
   "sebagai holdout yang gak disentuh sama sekali, baru dicek di akhir, buat buktiin sistemnya "
   "gak cuma disetel biar pas ke 1 rekaman doang.", 8)

qa("Isn't this dataset too clean / too easy compared to real factory conditions?",
   "No — it's a real physical experiment run to actual failure, so it has genuine sensor "
   "noise and unpredictability. We deliberately avoided cleaner benchmark datasets (like "
   "CWRU) as the main proof, specifically because those are too easy and every model scores "
   "near-perfect on them.",
   "Enggak — ini eksperimen fisik beneran sampai gagal beneran, jadi ada noise sensor asli dan "
   "gak bisa ditebak. Kita sengaja hindarin dataset benchmark yang lebih bersih (kayak CWRU) "
   "sebagai bukti utama, justru karena itu terlalu gampang, semua model dapet nilai nyaris "
   "sempurna di situ.", 9)

qa("Does the dataset include temperature, current, or pressure data?",
   "No, it's vibration-only. The physical test rig did have thermocouples pictured in the "
   "documentation, but that temperature data was never actually released as part of the "
   "public dataset.",
   "Enggak, cuma vibrasi doang. Alat uji fisiknya emang ada thermocouple di gambar "
   "dokumentasinya, tapi data suhu itu gak pernah dirilis jadi bagian dataset publiknya.", 10)

doc.add_page_break()

# ======================================================================
# C. Features
# ======================================================================
h2("Features (RMS, kurtosis, envelope, BPFO/BPFI/BSF/FTF)", "C")

qa("What features do you extract from the raw vibration signal?",
   "Time-domain: RMS, kurtosis, skewness, crest factor. Frequency-domain: envelope spectrum "
   "energy at three bearing fault frequencies — BPFO, BPFI, BSF.",
   "Time-domain: RMS, kurtosis, skewness, crest factor. Frequency-domain: energi envelope "
   "spectrum di tiga frekuensi fault bearing — BPFO, BPFI, BSF.", 11)

qa("What is RMS and why use it?",
   "Root Mean Square — the standard vibration-severity measurement, the same idea used in the "
   "ISO 10816 industry standard. Bigger vibration means bigger RMS.",
   "Root Mean Square — ukuran standar keparahan getaran, sama kayak yang dipakai standar "
   "industri ISO 10816. Getaran makin besar, RMS-nya makin besar.", 12)

qa("What is kurtosis and why does it matter for bearings?",
   "A statistical measure of how “spiky” a signal is. A healthy bearing vibrates smoothly "
   "(kurtosis around 3). A damaged bearing produces sharp metal-on-metal impacts, which push "
   "kurtosis much higher.",
   "Ukuran statistik seberapa \"runcing\" sinyalnya. Bearing sehat bergetar halus (kurtosis "
   "sekitar 3). Bearing rusak menghasilkan benturan logam-ke-logam tajam, yang bikin kurtosis "
   "naik jauh lebih tinggi.", 13)

qa("What is crest factor?",
   "Peak amplitude divided by RMS — another way to measure how “spiky” the signal is, more "
   "sensitive to just one extreme moment in the signal.",
   "Amplitudo puncak dibagi RMS — cara lain ngukur seberapa \"runcing\" sinyalnya, lebih "
   "sensitif ke satu momen ekstrem doang di sinyalnya.", 14)

qa("What is envelope analysis and why is it needed?",
   "A 3-step method — band-pass filter, then Hilbert transform, then FFT — used to reveal "
   "small periodic impacts that are otherwise hidden inside the much bigger structural "
   "vibration of the machine.",
   "Metode 3 langkah — band-pass filter, transformasi Hilbert, terus FFT — buat nunjukin "
   "benturan periodik kecil yang kalau enggak, ketutup sama getaran struktural mesin yang "
   "jauh lebih besar.", 15)

qa("What are BPFO, BPFI, BSF, FTF?",
   "Bearing fault frequencies, calculated from the bearing's physical geometry — number of "
   "rolling elements, diameters, contact angle, shaft speed. They stand for Outer race, Inner "
   "race, Ball spin, and cage (Fundamental Train) frequency.",
   "Frekuensi fault bearing, dihitung dari geometri fisik bearing-nya — jumlah rolling "
   "element, diameter, sudut kontak, kecepatan poros. Singkatan dari frekuensi outer race, "
   "inner race, ball spin, dan kecepatan sangkar (cage).", 16)

qa("How do you know these frequency calculations are correct?",
   "They were checked against the reference values already documented in the team's own "
   "methodology report, and they matched exactly — 236.4, 296.9, 139.9, and 14.8 Hz.",
   "Dicek ulang sama nilai referensi yang udah didokumentasikan di laporan metodologi tim "
   "sendiri, dan hasilnya cocok persis — 236,4, 296,9, 139,9, dan 14,8 Hz.", 17)

qa("Why not just use the raw signal, why go through all this feature engineering?",
   "A small bearing defect is tiny compared to all the other vibration in the machine — it's "
   "almost invisible in the raw signal. The envelope spectrum pulls that small signal out and "
   "shows it clearly, at its own known frequency.",
   "Cacat kecil di bearing itu kecil banget dibanding getaran lain di mesin — hampir gak "
   "keliatan di sinyal mentah. Envelope spectrum itu narik sinyal kecil itu keluar dan "
   "nunjukin jelas, di frekuensi khususnya sendiri.", 18)

doc.add_page_break()

# ======================================================================
# D. Anomaly Detection
# ======================================================================
h2("Anomaly Detection (Isolation Forest)", "D")

qa("What machine learning model do you use for anomaly detection?",
   "Isolation Forest — an unsupervised anomaly-detection algorithm.",
   "Isolation Forest — algoritma deteksi anomali yang unsupervised (tanpa label).", 19)

qa("Why unsupervised, why not a labeled/supervised model?",
   "Because the dataset has no per-timestamp labels — we only know how each run ended, not "
   "which exact moment counts as “already abnormal.” That's actually realistic: most real "
   "industrial data doesn't have hand-labeled anomaly tags either.",
   "Karena datasetnya gak punya label per-titik-waktu — kita cuma tau gimana tiap run "
   "berakhir, bukan momen persis mana yang dianggap \"udah abnormal\". Ini emang realistis: "
   "kebanyakan data industri beneran juga gak punya label anomali yang dibuat manual.", 20)

qa("How does Isolation Forest actually work?",
   "It splits the data randomly, over and over. Outliers get isolated in fewer splits than "
   "normal points, because they're easier to separate from everything else. Fewer splits "
   "needed means a higher anomaly score.",
   "Dia belah-belah data secara acak, berulang-ulang. Outlier itu keisolir dalam lebih "
   "sedikit belahan dibanding titik normal, karena dia lebih gampang dipisahin dari yang "
   "lain. Makin sedikit belahan yang dibutuhkan, makin tinggi skor anomalinya.", 21)

qa("What data is the model trained on?",
   "Only the first 50% of each run, assumed healthy — a conservative cut, more careful than "
   "the 60-70% suggested in early research, so no degrading data leaks into training.",
   "Cuma 50% awal tiap run, diasumsikan sehat — potongan yang konservatif, lebih hati-hati "
   "dari 60-70% yang disaranin riset awal, biar data yang udah mulai rusak gak bocor ke "
   "training.", 22)

qa("How is the alarm threshold chosen?",
   "Statistically — the healthy-period score's mean plus 3 standard deviations. Not an "
   "arbitrary hand-picked number.",
   "Secara statistik — rata-rata skor periode sehat ditambah 3 standar deviasi. Bukan angka "
   "sembarang yang dipilih tangan.", 23)

doc.add_page_break()

# ======================================================================
# E. Health Score & RUL
# ======================================================================
h2("Health Score & Remaining Useful Life (RUL)", "E")

qa("How is the Health Score (0-100) calculated?",
   "PCA combines the features that move consistently as damage grows into one number (the "
   "first principal component), smoothed over time, then scaled to 0-100 relative to that "
   "specific bearing's own healthy baseline and worst observed value.",
   "PCA menggabungkan fitur-fitur yang bergerak konsisten seiring kerusakan jadi 1 angka "
   "(komponen utama pertama), dihaluskan seiring waktu, terus diskalakan ke 0-100 relatif "
   "terhadap baseline sehat dan nilai terburuk bearing itu sendiri.", 24)

qa("Why PCA instead of manually weighting the features yourself?",
   "Manual weights are basically just guesses. PCA calculates the weighting mathematically "
   "from the real pattern in the data, so it's more defensible.",
   "Bobot manual itu intinya cuma tebakan. PCA ngitung bobotnya secara matematis dari pola "
   "asli di datanya, jadi lebih bisa dipertanggungjawabkan.", 25)

qa("Sometimes Health Score is still high but Alarm Status already says Warning or Critical — "
   "is that a bug?",
   "No — they're built for different jobs on purpose. Alarm Status reacts fast, catching "
   "small anomalies right away. Health Score is deliberately smoothed to show the bigger "
   "trend. That's by design.",
   "Bukan — mereka memang dibuat buat tugas beda dengan sengaja. Alarm Status bereaksi "
   "cepat, nangkep anomali kecil langsung. Health Score sengaja dihaluskan buat nunjukin tren "
   "besarnya. Itu memang desainnya.", 26)

qa("When does RUL start being calculated?",
   "Only after the first confirmed alarm. Before that, the system reports “monitoring, not "
   "stable yet” instead of guessing — guessing before real evidence exists wouldn't be "
   "statistically honest.",
   "Baru setelah alarm pertama terkonfirmasi. Sebelum itu, sistemnya bilang \"monitoring, "
   "belum stabil\" daripada nebak-nebak — nebak sebelum ada bukti nyata itu gak jujur secara "
   "statistik.", 27)

qa("What's the RUL model / formula?",
   "An exponential curve, HI(t) = p1 times e to the power of p2 times t, plus p3, fitted to "
   "the health trend after the alarm, refit as new data arrives. The end-of-life point is "
   "where that curve is projected to cross a failure threshold; RUL is that point minus now.",
   "Kurva eksponensial, HI(t) = p1 dikali e pangkat (p2 dikali t), ditambah p3, di-fit ke tren "
   "kesehatan setelah alarm, dihitung ulang tiap ada data baru. Titik akhir umur itu di mana "
   "kurva itu diperkirakan nyentuh ambang gagal; RUL itu titik itu dikurangi sekarang.", 28)

qa("Why exponential, why not a neural network or LSTM?",
   "There's very little failure data available — only a couple of full run-to-failure "
   "examples. A simple statistical model fits that amount of data better, carries less risk "
   "of overfitting, and is easier to explain to maintenance staff.",
   "Data kegagalan yang ada dikit banget — cuma beberapa contoh run-to-failure lengkap. Model "
   "statistik sederhana lebih cocok buat jumlah data segitu, resiko overfitting lebih kecil, "
   "dan lebih gampang dijelasin ke staf maintenance.", 29)

qa("Where does the failure threshold (ET) come from?",
   "From the OTHER run's own end-of-life Health Index value — never from the current run's "
   "own future data, so there's no leakage.",
   "Dari nilai Health Index di akhir umur run yang LAIN — gak pernah dari data masa depan run "
   "yang sekarang, jadi gak ada kebocoran data.", 30)

qa("How accurate is the RUL prediction, honestly?",
   "Moderate — RMSE around 26 hours on the development run and about 106 hours on the "
   "1073-hour holdout run, measured on the last 10% of life. It's a known, acknowledged "
   "limitation the team is actively improving, not something hidden.",
   "Sedang-sedang aja, jujurnya — RMSE sekitar 26 jam di run pengembangan dan sekitar 106 jam "
   "di run holdout 1073 jam, diukur di 10% terakhir umurnya. Ini keterbatasan yang diakui dan "
   "lagi diperbaiki, bukan yang disembunyiin.", 31)

qa("Why does the RUL number sometimes change a lot in a short time?",
   "The curve keeps refitting as new data comes in. Right after the alarm triggers, there's "
   "very little data, so the estimate can swing. It settles down as more of the real "
   "degradation trend becomes visible.",
   "Kurvanya terus dihitung ulang tiap ada data baru. Begitu alarm baru nyala, datanya masih "
   "dikit banget, jadi tebakannya bisa goyang. Dia jadi stabil begitu tren kerusakan aslinya "
   "makin keliatan.", 32)

doc.add_page_break()

# ======================================================================
# F. False-Alarm Reduction
# ======================================================================
h2("False-Alarm Reduction", "F")

qa("How does the system reduce false alarms compared to a simple threshold system?",
   "Two rules combined: Persistence — a single abnormal feature must hold for several "
   "consecutive readings before it counts. Voting — two or more features abnormal at the "
   "same time escalate immediately, since that's unlikely to be a coincidence.",
   "Dua aturan digabung: Persistence — 1 fitur abnormal harus bertahan beberapa pembacaan "
   "berturut-turut baru dihitung. Voting — 2 fitur atau lebih abnormal bersamaan langsung "
   "dieskalasi, karena itu kecil kemungkinan cuma kebetulan.", 33)

qa("Can you prove the false-alarm reduction actually works, with real numbers?",
   "Yes — on the holdout run, a single-threshold baseline produced 507 separate alarm "
   "episodes, versus just 1 after persistence and voting were applied, evaluated on the "
   "exact same recording.",
   "Bisa — di run holdout, baseline threshold tunggal menghasilkan 507 episode alarm "
   "terpisah, dibanding cuma 1 setelah persistence dan voting diterapkan, dievaluasi di "
   "rekaman yang sama persis.", 34)

qa("What's “burn-in” and why is it needed?",
   "The first 5% of a run is excluded from alarm evaluation — a calibration period so the "
   "system has enough healthy data to learn what “normal” looks like before judging anything.",
   "5% awal dari sebuah run dikecualikan dari evaluasi alarm — masa kalibrasi biar sistemnya "
   "punya cukup data sehat buat belajar seperti apa \"normal\" sebelum menilai apapun.", 35)

qa("What's hysteresis, why doesn't the alarm clear immediately when one reading looks normal "
   "again?",
   "Recovery requires several consecutive normal readings before the status downgrades — "
   "otherwise one noisy “back to normal” reading would chop a real alarm episode into many "
   "flickering ones.",
   "Recovery butuh beberapa pembacaan normal berturut-turut sebelum statusnya turun — kalau "
   "enggak, 1 pembacaan noise \"balik normal\" bakal motong 1 episode alarm asli jadi banyak "
   "yang kedip-kedip.", 36)

doc.add_page_break()

# ======================================================================
# G. Validation & This Week's Bug Fix
# ======================================================================
h2("Validation & This Week's Bug Fix", "G")

qa("How do you validate the system isn't just overfit to one recording?",
   "Test 2 is used for development. Test 3 — 1073.3 hours, 6.5 times longer — is kept "
   "completely untouched until final validation, using the exact same code and thresholds, "
   "just scaled proportionally to run length, not hand-tuned per run.",
   "Test 2 dipakai buat pengembangan. Test 3 — 1073,3 jam, 6,5 kali lebih panjang — disimpan "
   "gak disentuh sama sekali sampai validasi akhir, pakai kode dan ambang batas yang sama "
   "persis, cuma diskalakan proporsional ke panjang run, bukan disetel tangan per run.", 37)

qa("What went wrong this week, and how did you find it?",
   "Numbers on the long holdout run looked inconsistent with an earlier draft. Instead of "
   "assuming either was right, we investigated, and found a real bug: the alarm's voting "
   "rule could fire from a single noisy reading, no persistence needed. On the holdout run, "
   "that fired while the bearing was still 100% healthy, and it corrupted every RUL "
   "prediction that came after it.",
   "Angka di run holdout yang panjang keliatan gak konsisten sama draft sebelumnya. "
   "Daripada asumsiin salah satu yang benar, kita investigasi, dan nemu bug beneran: aturan "
   "voting alarm bisa bunyi dari 1 pembacaan noise doang, tanpa perlu persistence. Di run "
   "holdout, itu bunyi pas bearing-nya masih 100% sehat, dan itu ngerusak semua prediksi RUL "
   "setelahnya.", 38)

qa("What was the fix, and what actually changed?",
   "Voting now needs two consecutive abnormal-together readings before it counts as a real "
   "alarm. Result: RMSE improved from 475 hours to 106 hours (4.5 times better) on the "
   "holdout run, and the false trigger point moved from a meaningless noise spike to a point "
   "of real, multi-feature degradation.",
   "Voting sekarang butuh 2 pembacaan abnormal-bersamaan berturut-turut baru dianggap alarm "
   "beneran. Hasilnya: RMSE membaik dari 475 jam jadi 106 jam (4,5 kali lebih baik) di run "
   "holdout, dan titik trigger yang salah pindah dari lonjakan noise gak berarti ke titik "
   "degradasi multi-fitur yang beneran.", 39)

doc.add_page_break()

# ======================================================================
# H. Limitations & Honesty
# ======================================================================
h2("Limitations & Honesty", "H")

qa("What does this prototype NOT do yet?",
   "Vibration only, no temperature/current/pressure yet. Tested on one constant operating "
   "condition, while real factories vary speed and load. The single exponential RUL model "
   "can be unstable soon after triggering. And the failure threshold currently comes from "
   "just one other run, not a distribution across many.",
   "Vibrasi doang, belum ada suhu/arus/tekanan. Ditest di 1 kondisi operasi konstan, padahal "
   "pabrik beneran variasi kecepatan dan beban. Model RUL eksponensial tunggal bisa gak "
   "stabil segera setelah trigger. Dan ambang batas gagal sekarang cuma dari 1 run lain, "
   "bukan distribusi dari banyak run.", 40)

qa("Is the “likely root cause” text on the dashboard a real independent AI diagnosis?",
   "No, and the team is upfront about that. It uses a known historical fact — this specific "
   "bearing's documented failure mode — rather than doing a fully independent real-time "
   "diagnosis. It's an honest, disclosed limitation, not hidden.",
   "Enggak, dan tim terbuka soal itu. Dia pakai fakta historis yang udah diketahui — mode "
   "kegagalan bearing spesifik ini yang udah didokumentasikan — bukan diagnosis mandiri "
   "real-time yang beneran. Ini keterbatasan yang jujur diungkap, bukan disembunyikan.", 41)

qa("Would this still work if a different type of defect happened (inner race, ball, not "
   "outer race)?",
   "Yes in principle — the same envelope-analysis method would look at the BPFI or BSF line "
   "instead of BPFO. The current root-cause text just isn't automatically picking the right "
   "one yet — it's based on the known documented failure mode per run, not a live comparison.",
   "Iya secara prinsip — metode envelope-analysis yang sama akan liat garis BPFI atau BSF, "
   "bukan BPFO. Teks root-cause yang sekarang belum otomatis milih yang tepat — dia berbasis "
   "mode kegagalan yang udah diketahui per run, bukan perbandingan live.", 42)

doc.add_page_break()

# ======================================================================
# I. The Dashboard / Prototype
# ======================================================================
h2("The Dashboard / Prototype Itself", "I")

qa("Is this dashboard connected to a real machine right now?",
   "No — it replays historical recorded data, snapshot by snapshot, to simulate what "
   "real-time monitoring would look like. Live sensor streaming is future work.",
   "Enggak — dia replay data rekaman historis, snapshot demi snapshot, buat simulasiin "
   "gimana rasanya monitoring real-time. Streaming sensor live itu kerjaan masa depan.", 43)

qa("What are the three tabs and what does each one show?",
   "Trend shows health score and anomaly score over time. Signal Detail shows the raw "
   "vibration and the envelope spectrum “fingerprint.” Alerts & Recommendation shows the "
   "reason, likely root cause, and recommended action.",
   "Trend nunjukin health score dan anomaly score dari waktu ke waktu. Signal Detail "
   "nunjukin vibrasi mentah dan \"fingerprint\" envelope spectrum. Alerts & Recommendation "
   "nunjukin alasan, kemungkinan penyebab, dan rekomendasi tindakan.", 44)

qa("What technology is the dashboard built with?",
   "Python, Streamlit for the web app itself, Plotly for the charts, scikit-learn for "
   "Isolation Forest and PCA, and SciPy for the signal processing.",
   "Python, Streamlit buat aplikasi webnya, Plotly buat grafiknya, scikit-learn buat "
   "Isolation Forest dan PCA, dan SciPy buat pemrosesan sinyalnya.", 45)

qa("Can this scale to monitor many machines at once?",
   "Architecturally yes — each bearing or machine gets its own model and thresholds. The "
   "current prototype demos one machine at a time for clarity, not because of a technical "
   "limit.",
   "Secara arsitektur bisa — tiap bearing atau mesin punya model dan ambang batasnya "
   "sendiri. Prototipe sekarang demo 1 mesin dalam satu waktu biar jelas, bukan karena "
   "keterbatasan teknis.", 46)

doc.add_page_break()

# ======================================================================
# J. Roadmap
# ======================================================================
h2("Roadmap & Future Work", "J")

qa("What's the plan once Astra provides real sensor data?",
   "The feature-extraction and alarm modules are built to be generic — new sensor types like "
   "temperature, current, and pressure would just become new feature columns, without "
   "needing to redesign the anomaly detection, alarm, or RUL logic.",
   "Modul ekstraksi fitur dan alarm dibuat generik — jenis sensor baru kayak suhu, arus, dan "
   "tekanan tinggal jadi kolom fitur baru, tanpa perlu ngerombak ulang logic deteksi anomali, "
   "alarm, atau RUL.", 47)

qa("What are the next steps after this week?",
   "Stabilize the early-stage RUL prediction, prepare for multi-sensor data, test the same "
   "approach on other bearings, and start designing for live-streaming data instead of "
   "replaying historical files.",
   "Stabilin prediksi RUL di tahap awal, siapin buat data multi-sensor, tes pendekatan yang "
   "sama di bearing lain, dan mulai desain buat data streaming live, bukan replay file "
   "historis.", 48)

qa("What's the long-term rollout plan at Astra?",
   "Phase 1: instrument 3-5 real motors and run in shadow mode alongside technicians. "
   "Phase 2: live PLC/SCADA connectivity through OPC UA or Modbus. Phase 3: scale across "
   "factories and extend fault coverage to stator and rotor-bar issues too.",
   "Fase 1: pasang sensor di 3-5 motor beneran, jalan shadow mode bareng teknisi. Fase 2: "
   "koneksi live PLC/SCADA lewat OPC UA atau Modbus. Fase 3: skalain ke banyak pabrik dan "
   "perluas cakupan fault ke masalah stator dan rotor-bar juga.", 49)

doc.add_page_break()

# ======================================================================
# K. Team & Process
# ======================================================================
h2("Team & Process", "K")

qa("How did the team divide the work?",
   "Three people, working on one shared codebase — one focused on the physics/methodology "
   "and validation report, and the dashboard presenters split the live demo into three parts "
   "covering different sections of the same Streamlit app.",
   "Tiga orang, kerja di 1 codebase yang sama — satu fokus ke fisika/metodologi dan laporan "
   "validasi, dan presenter dashboard-nya bagi demo live jadi 3 bagian yang nutupin section "
   "beda-beda dari aplikasi Streamlit yang sama.", 50)

qa("How long has the team been working on this?",
   "Since early July 2026, following the locked methodology blueprint, built out actively "
   "over the following weeks with weekly progress checkpoints.",
   "Sejak awal Juli 2026, ngikutin blueprint metodologi yang udah dikunci, dikembangin aktif "
   "selama minggu-minggu berikutnya dengan checkpoint progress mingguan.", 51)

doc.add_page_break()

# ======================================================================
# L. Tricky / Conceptual
# ======================================================================
h2("Tricky / Conceptual Questions", "L")

qa("If Health Score looks fine but Alarm Status already says Critical, which one should I "
   "trust?",
   "Trust Alarm Status for immediate action — it's the faster, more sensitive signal. Health "
   "Score is the smoothed trend view, good for the bigger picture, not moment-to-moment "
   "decisions.",
   "Percaya Alarm Status buat tindakan langsung — itu sinyal yang lebih cepat dan sensitif. "
   "Health Score itu tampilan tren yang dihaluskan, bagus buat gambaran besar, bukan "
   "keputusan sesaat.", 52)

qa("Could the false-alarm-reduction rules also delay a REAL alarm and cause harm?",
   "Yes, that's a real tradeoff, and the team measured it rather than ignoring it — "
   "persistence adds a small delay in exchange for far fewer false alarms, and that delay is "
   "reported alongside the false-alarm-reduction numbers, not hidden.",
   "Iya, itu tradeoff yang nyata, dan tim ngukurnya, bukan diabaikan — persistence nambah "
   "sedikit delay demi false alarm yang jauh lebih sedikit, dan delay itu dilaporin "
   "bareng-bareng sama angka pengurangan false alarm, gak disembunyikan.", 53)

qa("What happens before the “knee point” — before enough data exists to trust a prediction?",
   "The system deliberately reports “monitoring, not stable yet” instead of guessing — "
   "guessing before real evidence exists wouldn't be statistically honest.",
   "Sistemnya sengaja bilang \"monitoring, belum stabil\" daripada nebak — nebak sebelum ada "
   "bukti nyata itu gak jujur secara statistik.", 54)

qa("Why not just use deep learning for everything — detection AND RUL?",
   "There's too little failure data available for deep learning to be reliable. Simpler "
   "statistical and machine-learning methods are more appropriate given the data volume, and "
   "easier to explain and trust for non-technical maintenance staff.",
   "Data kegagalan yang ada terlalu sedikit buat deep learning bisa diandalkan. Metode "
   "statistik dan machine learning yang lebih sederhana lebih cocok buat volume data segini, "
   "dan lebih gampang dijelasin dan dipercaya buat staf maintenance yang non-teknis.", 55)

qa("How is this different from just setting a fixed alarm threshold on each sensor?",
   "A fixed threshold treats every sensor independently and reacts to any single blip. This "
   "system looks at multiple features together, requires either sustained abnormality or "
   "several features agreeing at once, and adds a time-based prediction on top — that's the "
   "core difference the whole project is built around.",
   "Threshold tetap itu memperlakukan tiap sensor sendiri-sendiri dan bereaksi ke kedipan "
   "sesaat manapun. Sistem ini liat beberapa fitur bareng-bareng, butuh abnormalitas yang "
   "bertahan atau beberapa fitur sepakat bersamaan, dan nambahin prediksi berbasis waktu di "
   "atasnya — itu perbedaan inti yang jadi dasar seluruh project ini.", 56)

OUT_PATH.parent.mkdir(exist_ok=True)
doc.save(str(OUT_PATH))
print("Saved:", OUT_PATH)
print("Total questions:", 56)
