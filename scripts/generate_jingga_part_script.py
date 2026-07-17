"""Generate Jingga's solo demo script (Part 3 of 3: Trend tab, Signal Detail
tab, Alerts & Recommendation tab, Closing) as a standalone docx.

v5: SHORT, live-action-first. Every line is action + speech together (move
something, then say what changed and why it matters). No long theory. A
separate "backup notes" section at the end holds the deeper explanations,
clearly marked optional, so the main script stays short enough to actually
remember.

Usage (from project root):
    python -m scripts.generate_jingga_part_script

Produces docs/Jingga_Part3_Demo_Script.docx
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "docs" / "Jingga_Part3_Demo_Script_v7.docx"

NAVY = RGBColor(0x0B, 0x2A, 0x4A)
ACCENT = RGBColor(0x03, 0x69, 0xA1)
GRAY = RGBColor(0x55, 0x5B, 0x66)
SPOKEN = RGBColor(0x11, 0x14, 0x18)
ACTION_COLOR = RGBColor(0xB4, 0x5C, 0x00)
ID_COLOR = RGBColor(0x0F, 0x76, 0x4A)
NOTE_COLOR = RGBColor(0x8A, 0x3A, 0x3A)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(12)

for section in doc.sections:
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)


def h1(text, color=NAVY, size=22):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.bold = True; r.font.color.rgb = color
    p.space_after = Pt(8)
    return p


def h2(text, color=ACCENT):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(15); r.font.bold = True; r.font.color.rgb = color
    p.space_before = Pt(16); p.space_after = Pt(6)
    return p


def step(action_text, say_en, say_id, size=13):
    """One live-action beat: orange action line immediately followed by the
    black spoken line (in quotes) and the green Indonesian meaning under it.
    This is the core unit of the whole script — do this, then say this."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(action_text)
    r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = ACTION_COLOR

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    p2.paragraph_format.line_spacing = 1.2
    r2 = p2.add_run("“" + say_en + "”")
    r2.font.size = Pt(size); r2.font.color.rgb = SPOKEN

    p3 = doc.add_paragraph()
    p3.paragraph_format.left_indent = Cm(0.5)
    p3.paragraph_format.space_after = Pt(4)
    r3 = p3.add_run(say_id)
    r3.font.size = Pt(10.5); r3.font.italic = True; r3.font.color.rgb = ID_COLOR
    return p


def say_only(say_en, say_id, size=13):
    """A spoken line with no physical action attached to it."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run("“" + say_en + "”")
    r.font.size = Pt(size); r.font.color.rgb = SPOKEN
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.5)
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run(say_id)
    r2.font.size = Pt(10.5); r2.font.italic = True; r2.font.color.rgb = ID_COLOR
    return p


def backup_qa(question, answer_en, answer_id):
    p = doc.add_paragraph()
    r = p.add_run("If asked: " + question)
    r.font.bold = True; r.font.size = Pt(12); r.font.color.rgb = NAVY
    p.space_before = Pt(12); p.space_after = Pt(3)
    p2 = doc.add_paragraph()
    r2 = p2.add_run("“" + answer_en + "”")
    r2.font.size = Pt(11.5); r2.font.color.rgb = SPOKEN
    p2.space_after = Pt(2)
    p3 = doc.add_paragraph()
    p3.paragraph_format.left_indent = Cm(0.5)
    r3 = p3.add_run(answer_id)
    r3.font.size = Pt(10); r3.font.italic = True; r3.font.color.rgb = ID_COLOR
    p3.space_after = Pt(4)


def hr():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), "1F6FC4")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.space_after = Pt(4)


def small_note(text, color=GRAY):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10.5); r.font.italic = True; r.font.color.rgb = color
    p.space_after = Pt(8)


# ======================================================================
# Cover
# ======================================================================
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("My Part — Live Demo Script")
r.font.size = Pt(26); r.font.bold = True; r.font.color.rgb = NAVY

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Jingga · Part 3 of 3 · SHORT version")
r.font.size = Pt(15); r.font.color.rgb = ACCENT

sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("Trend tab → Signal Detail tab → Alerts & Recommendation tab → Closing")
r.font.size = Pt(11); r.font.color.rgb = GRAY

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("How this works: ")
r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = NAVY
r2 = p.add_run(
    "orange = what you physically do (click, drag). Black in quotes = what "
    "you say, right after you do it. Green italic = just the Indonesian "
    "meaning, for you only, never said out loud. This is the SHORT version "
    "— memorize this part. At the very end there's a separate “backup” "
    "section, only for if Miss asks something extra — you don't need to "
    "memorize that part, just read it once so you're not surprised."
)
r2.font.size = Pt(11); r2.font.color.rgb = GRAY
p.space_after = Pt(4)

p = doc.add_paragraph()
r = p.add_run(
    "Cara pakai: oranye = aksi fisik (klik, geser). Hitam pakai kutip = yang "
    "diucapin, langsung setelah aksinya. Hijau miring = artinya doang buat "
    "kamu, gak diucapin. Ini versi PENDEK — hafalin bagian ini aja. Di "
    "paling bawah ada bagian “cadangan”, cuma buat kalau-kalau Miss nanya "
    "lebih — gak perlu dihafal, baca sekali aja biar gak kaget."
)
r.font.size = Pt(10.5); r.font.italic = True; r.font.color.rgb = ID_COLOR
p.space_after = Pt(10)

p = doc.add_paragraph()
r = p.add_run("Hand-off: ")
r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = NAVY
r2 = p.add_run("Syakira finishes with something like “...now Jingga will show you "
               "the history and the proof.” That's your cue.")
r2.font.size = Pt(11); r2.font.color.rgb = GRAY

doc.add_page_break()

# ======================================================================
# TREND TAB — short, live
# ======================================================================
h1("Part A — “Trend” Tab")
small_note("At this point: slider around 762, health score about 83, status Critical. "
           "You don't need to touch anything here — just talk over what's already on screen.")

step("ACTION — point at the left chart (health score).",
     "So, this chart, it show the whole story — health score, start at 100, going down to now. "
     "This orange line here, that's the exact moment the alarm turn on. And you can see, the "
     "score already start dropping before it get critical — that's the whole point, we catch "
     "it early, not after it already broken.",
     "Jadi, grafik ini nunjukin cerita lengkapnya — health score, mulai dari 100, turun sampai "
     "sekarang. Garis oranye ini, itu momen persis alarm nyala. Dan keliatan, skornya udah mulai "
     "turun sebelum jadi critical — itu intinya, kita nangkep dari awal, bukan setelah rusak.")

step("ACTION — point at the right chart (anomaly score).",
     "This chart, it's simple — dark line is how strange the signal look, red line is the "
     "limit. The moment dark line cross the red line, that's the trigger, that's when alarm "
     "turn on. We don't pick this red line randomly — it's a standard limit from statistic, so "
     "we know it's a real signal, not just noise.",
     "Grafik ini simpel — garis gelap itu seberapa aneh sinyalnya, garis merah itu batasnya. "
     "Begitu garis gelap ngelewatin garis merah, itu triggernya, itu pas alarm nyala. Kita gak "
     "asal pilih garis merah ini — itu batas standar dari statistik, jadi kita tau itu sinyal "
     "beneran, bukan cuma noise.")

doc.add_page_break()

# ======================================================================
# SIGNAL DETAIL TAB — short, live
# ======================================================================
h1("Part B — “Signal Detail” Tab")

step("ACTION — click “Signal Detail.” Remember the slider number first (around 762) — "
     "you'll come back to it.",
     "This tab, it's the real proof, straight from the sensor.",
     "Tab ini, bukti aslinya, langsung dari sensor.")

h2("Left chart — Raw vibration snapshot")
step("ACTION — point at the left chart.",
     "This left chart, it's the raw signal, straight from the sensor, twenty thousand readings "
     "every second. Right now the machine already Critical, and this signal is already a bit "
     "busier than a healthy one — around one and a half times bigger. It's not the most extreme "
     "point yet, but you can already see it's not clean and calm anymore.",
     "Grafik kiri ini, sinyal mentah, langsung dari sensor, dua puluh ribu bacaan tiap detik. "
     "Sekarang mesinnya udah Critical, dan sinyal ini udah agak lebih ramai dibanding pas sehat "
     "— sekitar satu setengah kali lebih besar. Ini belum titik paling parah, tapi udah keliatan "
     "gak bersih dan tenang lagi.")

h2("Right chart — Envelope spectrum")
step("ACTION — drag the slider all the way back to 0.",
     "Watch — I move back to the start, healthy condition. Look at this chart on the right, "
     "flat, quiet, nothing happening.",
     "Liat — aku geser balik ke awal, kondisi sehat. Liat grafik kanan ini, datar, tenang, gak "
     "ada apa-apa.")

step("ACTION — drag the slider forward again, back to around 762.",
     "Now watch — I move close to failure. See this spike appear, right here? This is not "
     "random. This spike, it match the exact frequency this bearing produce if the outer ring "
     "is damaged. That's real physical proof, not just a statistic guess.",
     "Sekarang liat — aku geser deket kegagalan. Liat lonjakan ini muncul, di sini? Ini bukan "
     "acak. Lonjakan ini, cocok persis sama frekuensi yang dihasilkan bearing ini kalau cincin "
     "luarnya rusak. Itu bukti fisik nyata, bukan cuma tebakan statistik.")

doc.add_page_break()

# ======================================================================
# ALERTS & RECOMMENDATION TAB — short
# ======================================================================
h1("Part C — “Alerts & Recommendation” Tab")

step("ACTION — click “Alerts & Recommendation.”",
     "This tab answer the last question — so what do we do? It says here, outer-race defect, "
     "and it tell us: schedule replacement now, don't wait until it break. This is what make "
     "it useful for a real maintenance team, not just numbers on a screen.",
     "Tab ini jawab pertanyaan terakhir — jadi kita harus ngapain? Di sini tertulis, cacat "
     "outer-race, dan dikasih tau: jadwalin penggantian sekarang, jangan tunggu sampai rusak. "
     "Ini yang bikin sistem ini berguna buat tim maintenance beneran, bukan cuma angka di layar.")

say_only("One honest thing — this root cause, it come from our own record of this exact "
         "bearing. The system is not guessing on its own, it's using what we already know.",
         "Satu hal jujur — root cause ini, itu dari catatan kita sendiri soal bearing ini. "
         "Sistemnya gak nebak sendirian, dia pakai yang udah kita tau.")

doc.add_page_break()

# ======================================================================
# CLOSING — one line, as requested
# ======================================================================
h1("Part D — Closing")
step("ACTION — look at the audience.",
     "Ok, that's all from us. Thank you.",
     "Oke, itu semua dari kami. Terima kasih.")

doc.add_page_break()

# ======================================================================
# BACKUP NOTES — optional depth, not required to memorize
# ======================================================================
h1("Backup Notes — only if Miss asks more", color=NOTE_COLOR, size=20)
small_note("Don't memorize this. Just read it once so nothing here surprises you. "
           "Everything above (Part A-D) is enough for the actual demo.")

h2("About the “Predicted trajectory” line (Trend tab)")
backup_qa("what's that other dashed line near the alarm?",
          "That's the system's own prediction — its best guess of where the health score is "
          "heading next, based on the trend so far. It only shows up once there's enough data "
          "after the alarm — before that, the system says it's still not sure enough to guess.",
          "Itu tebakan sistemnya sendiri — perkiraan terbaik ke arah mana health score bakal "
          "menuju, berdasarkan tren yang ada. Ini baru muncul setelah ada cukup data setelah "
          "alarm — sebelum itu, sistemnya bilang belum cukup yakin buat nebak.")

h2("About the taller, unlabeled spike (Signal Detail tab)")
backup_qa("there's an even taller spike further right, with no red line — what's that?",
          "That's a harmonic — basically a multiple of the same BPFO frequency we just talked "
          "about. It's actually more evidence for the same story, not a different problem.",
          "Itu harmonic — intinya kelipatan dari frekuensi BPFO yang sama tadi. Itu sebenernya "
          "bukti tambahan buat cerita yang sama, bukan masalah lain.")

h2("About the other frequency lines (BPFI, BSF, FTF)")
backup_qa("what are the other red lines for?",
          "Same idea as BPFO, just different parts of the bearing. BPFI at 296.9 Hertz is the "
          "inner ring, BSF at 139.9 is the rolling ball, FTF at 14.8 is the cage speed. All "
          "calculated from this bearing's real size, not guessed.",
          "Sama kayak BPFO, cuma beda bagian bearing-nya. BPFI di 296,9 Hertz itu cincin dalam, "
          "BSF di 139,9 itu bola-nya, FTF di 14,8 itu kecepatan sangkarnya. Semua dihitung dari "
          "ukuran asli bearing ini, bukan tebakan.")

h2("About the 3-sigma limit (Trend tab, right chart)")
backup_qa("why 3-sigma specifically, why not some other number?",
          "It's a common rule in statistics — normal readings almost always stay inside that "
          "range, about 99.7% of the time. So anything past it is very unlikely to be random "
          "noise.",
          "Itu aturan umum di statistik — bacaan normal hampir selalu ada di dalam rentang itu, "
          "sekitar 99,7% dari waktu. Jadi kalau udah lewat itu, kemungkinan besar bukan cuma "
          "noise acak.")

h2("About the root cause not being “independent AI diagnosis”")
backup_qa("does the system really diagnose the problem by itself?",
          "Not fully — if there's an alarm on this specific bearing, and we already know from "
          "our own records that this bearing ended up failing at the outer race, the dashboard "
          "shows that known fact. It's honest, not a fully independent real-time diagnosis yet.",
          "Enggak sepenuhnya — kalau ada alarm di bearing spesifik ini, dan kita udah tau dari "
          "catatan sendiri kalau bearing ini emang gagal di outer race, dashboard-nya nunjukin "
          "fakta yang udah diketahui itu. Ini jujur, belum diagnosis mandiri yang beneran "
          "real-time.")

h2("Is this real-time, or just replaying old data?")
backup_qa("is this actually live, or just a recording?",
          "Right now it's a replay of a real recording, from a real bearing that ran until it "
          "failed. We do this so we can test and prove the system honestly, on data we already "
          "know the ending of. Making it work on live, real-time data is one of our next steps.",
          "Sekarang ini replay dari rekaman asli, dari bearing beneran yang jalan sampai gagal. "
          "Kita lakuin gini biar bisa tes dan buktiin sistemnya secara jujur, pakai data yang "
          "udah kita tau hasil akhirnya. Bikin ini jalan di data live/real-time itu salah satu "
          "rencana kita selanjutnya.")

h2("Why look at the frequency spectrum at all — why not just the raw signal?")
backup_qa("why do you need the envelope spectrum, isn't the raw signal enough?",
          "Because a small bearing defect is tiny compared to all the other vibration in the "
          "machine — in the raw signal, it's almost hidden. The envelope spectrum is a way to "
          "pull that small signal out and see it clearly, at its own specific frequency.",
          "Karena cacat kecil di bearing itu kecil banget dibanding getaran lain di mesin — di "
          "sinyal mentah, itu hampir gak keliatan. Envelope spectrum itu cara buat narik sinyal "
          "kecil itu keluar dan liat dia jelas, di frekuensi khususnya sendiri.")

h2("What if it's a different kind of damage, not outer-race?")
backup_qa("would this still work for a different type of bearing damage?",
          "Yes — the same method works for inner-race or ball damage too, we'd just look at the "
          "BPFI or BSF line instead of BPFO. We picked outer-race for this demo because that's "
          "the known failure in this specific recording.",
          "Iya — metode yang sama jalan juga buat kerusakan di cincin dalam atau bola, kita "
          "tinggal liat garis BPFI atau BSF-nya, bukan BPFO. Kita pilih outer-race buat demo ini "
          "karena itu kegagalan yang udah diketahui di rekaman ini.")

h2("How confident are we in this — could it be wrong?")
backup_qa("how sure are you this isn't a false alarm?",
          "We can't be 100% sure with any system, but this alarm needs several abnormal "
          "readings together, not just one — that's exactly to avoid false alarms. And on this "
          "recording, the spike lines up with real bearing physics, not just statistics alone.",
          "Kita gak bisa 100% yakin di sistem manapun, tapi alarm ini butuh beberapa bacaan "
          "abnormal bareng-bareng, bukan cuma satu — itu emang buat ngehindarin false alarm. "
          "Dan di rekaman ini, lonjakannya cocok sama fisika bearing beneran, bukan cuma "
          "statistik doang.")

OUT_PATH.parent.mkdir(exist_ok=True)
doc.save(str(OUT_PATH))
print("Saved:", OUT_PATH)
